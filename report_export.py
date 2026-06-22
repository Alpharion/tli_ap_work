"""
report_export.py — PDF, DOCX and Excel export route handlers.

Registered in app.py via:
    from report_export import report_bp
    app.register_blueprint(report_bp)
"""

import io
import re

from flask import Blueprint, Response, jsonify, request, send_file

from excel_export import build_bulk_workbook

report_bp = Blueprint("report_export", __name__)


@report_bp.route("/api/export/excel", methods=["POST"])
def export_excel():
    try:
        data          = request.get_json(force=True)
        supply_chain  = data.get("supply_chain", {})
        product_input = data.get("product_input", supply_chain.get("product", {}).get("product_name", "export"))
        depth         = int(data.get("depth", 3))

        clean        = re.sub(r'[\\/*?\[\]:]', '', product_input).strip()[:26]
        sheet_prefix = f"1_{clean}"
        results      = [{"product_input": product_input, "supply_chain": supply_chain, "sheet_prefix": sheet_prefix}]

        wb_bytes = build_bulk_workbook(results, supply_chain.get("provider", "Frontend Export"), depth)
        buf      = io.BytesIO(wb_bytes)
        return send_file(buf, as_attachment=True,
                         download_name=f"{clean}_supply_chain.xlsx",
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@report_bp.route("/api/export/docx", methods=["POST"])
def export_docx():
    """Generate Word document using python-docx."""
    try:
        import datetime
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        product = data.get("product", {})
        tiers   = data.get("tiers", {})
        summary = data.get("summary", "")

        TIER_COLORS = [
            RGBColor(0x2B, 0x7F, 0xCC), RGBColor(0xD4, 0xAC, 0x0D),
            RGBColor(0x7D, 0x3C, 0x98), RGBColor(0xCA, 0x6F, 0x1E),
        ]
        TIER_HEX = ["2B7FCC", "D4AC0D", "7D3C98", "CA6F1E"]

        def set_cell_bg(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        def add_cell_text(cell, text, bold=False, size=10, color=None, italic=False):
            para = cell.paragraphs[0]
            run  = para.add_run(str(text if text is not None else "—"))
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        title = doc.add_heading("Supply Chain Intelligence Report", 0)
        title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        doc.add_paragraph(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}").runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()

        doc.add_heading("1. Product Overview", level=1)
        prod_rows = [
            ("Product Name",     product.get("product_name", "")),
            ("Category",         product.get("product_category", "")),
            ("Industry",         product.get("industry", "")),
            ("OEM Manufacturer", product.get("oem_manufacturer", "")),
            ("OEM Country",      product.get("oem_country", "")),
            ("Key Components",   ", ".join(product.get("key_components", []) or [])),
            ("Description",      product.get("description", "")),
        ]
        tbl     = doc.add_table(rows=len(prod_rows), cols=2)
        tbl.style = "Table Grid"
        col_w   = [Inches(2), Inches(4.5)]
        for i, (label, value) in enumerate(prod_rows):
            row = tbl.rows[i]
            row.cells[0].width = col_w[0]
            row.cells[1].width = col_w[1]
            add_cell_text(row.cells[0], label, bold=True, size=10)
            add_cell_text(row.cells[1], value, size=10)
            bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
            set_cell_bg(row.cells[0], bg)
            set_cell_bg(row.cells[1], bg)
        doc.add_paragraph()

        oems = data.get("oems", [])
        if oems:
            doc.add_heading("2. OEM Manufacturers", level=1)
            oem_headers = ["Rank", "Company", "Country", "Role", "Mkt Share", "Share %", "Source", "Notes"]
            owidths     = [Inches(0.5), Inches(1.2), Inches(0.8), Inches(1.0), Inches(0.8), Inches(0.7), Inches(1.0), Inches(1.5)]
            otbl        = doc.add_table(rows=len(oems) + 1, cols=8)
            otbl.style  = "Table Grid"
            hrow        = otbl.rows[0]
            for ci, (hdr_text, width) in enumerate(zip(oem_headers, owidths)):
                hrow.cells[ci].width = width
                add_cell_text(hrow.cells[ci], hdr_text, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
                set_cell_bg(hrow.cells[ci], "1A6B3C")
            for ri, oem in enumerate(oems, 1):
                row  = otbl.rows[ri]
                vals = [
                    str(oem.get("market_rank", "") or "—"), oem.get("company_name", "—"),
                    oem.get("country", "—"), oem.get("role", "—"),
                    oem.get("market_share", "—"), oem.get("market_share_pct", "—"),
                    oem.get("market_share_source", "—"), oem.get("notes", "—"),
                ]
                bg = "F0F8F0" if ri % 2 == 0 else "FFFFFF"
                for ci, (val, width) in enumerate(zip(vals, owidths)):
                    row.cells[ci].width = width
                    add_cell_text(row.cells[ci], val, bold=(ci == 1), size=9)
                    set_cell_bg(row.cells[ci], bg)
            doc.add_paragraph()

        if summary:
            doc.add_heading("3. Executive Summary", level=1)
            p = doc.add_paragraph(summary)
            p.runs[0].font.size = Pt(11)
            doc.add_paragraph()

        doc.add_heading("4. Supplier Tiers", level=1)
        tier_names = ["Tier 1", "Tier 2", "Tier 3", "Tier 4"]
        for ti, (tier_key, suppliers) in enumerate(tiers.items()):
            color  = TIER_COLORS[ti] if ti < len(TIER_COLORS) else RGBColor(0x80, 0x80, 0x80)
            hex_c  = TIER_HEX[ti]    if ti < len(TIER_HEX)    else "808080"
            tname  = tier_names[ti]  if ti < len(tier_names)  else f"Tier {ti+1}"
            h = doc.add_heading(f"{tname} Suppliers ({len(suppliers)})", level=2)
            h.runs[0].font.color.rgb = color
            if not suppliers:
                doc.add_paragraph("No suppliers identified.")
                continue
            headers = ["Company", "Country", "Supplies To", "Components Supplied", "Confidence"]
            tbl2   = doc.add_table(rows=len(suppliers) + 1, cols=5)
            tbl2.style = "Table Grid"
            widths = [Inches(1.6), Inches(1.0), Inches(1.4), Inches(1.9), Inches(0.7)]
            hrow   = tbl2.rows[0]
            for ci, (hdr, w) in enumerate(zip(headers, widths)):
                hrow.cells[ci].width = w
                add_cell_text(hrow.cells[ci], hdr, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
                set_cell_bg(hrow.cells[ci], hex_c)
            for ri, s in enumerate(suppliers):
                drow = tbl2.rows[ri + 1]
                vals = [
                    s.get("company_name", "—"), s.get("country", "—"),
                    s.get("supplies_to", "—"),
                    ", ".join((s.get("components_supplied") or [])[:3]),
                    s.get("confidence", "—"),
                ]
                bg = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
                for ci, (val, w) in enumerate(zip(vals, widths)):
                    drow.cells[ci].width = w
                    add_cell_text(drow.cells[ci], val, bold=(ci == 0), size=9)
                    set_cell_bg(drow.cells[ci], bg)
            doc.add_paragraph()

        doc.add_heading("5. Statistics", level=1)
        total     = sum(len(v) for v in tiers.values())
        countries = len(set(s.get("country", "") for arr in tiers.values() for s in arr if s.get("country")))
        stats = [
            ("Total Suppliers Identified", str(total)),
            ("Countries Represented",      str(countries)),
            ("Tier Depth Mapped",          str(len(tiers))),
            ("Industry",                   product.get("industry", "—")),
        ]
        stbl = doc.add_table(rows=len(stats), cols=2)
        stbl.style = "Table Grid"
        for i, (label, value) in enumerate(stats):
            stbl.rows[i].cells[0].width = Inches(2.5)
            stbl.rows[i].cells[1].width = Inches(4)
            add_cell_text(stbl.rows[i].cells[0], label, bold=True, size=10)
            add_cell_text(stbl.rows[i].cells[1], value, size=10)
            bg = "F0F4F8" if i % 2 == 0 else "FFFFFF"
            set_cell_bg(stbl.rows[i].cells[0], bg)
            set_cell_bg(stbl.rows[i].cells[1], bg)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", product.get("product_name", "supply_chain"))[:40]
        return Response(
            buf.read(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_supply_chain.docx"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@report_bp.route("/api/export/pdf", methods=["POST"])
def export_pdf():
    """Generate PDF from supply chain data using reportlab."""
    try:
        import datetime
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        product = data.get("product", {})
        tiers   = data.get("tiers", {})
        summary = data.get("summary", "")
        oems    = data.get("oems", [])

        buffer = io.BytesIO()
        doc    = SimpleDocTemplate(buffer, pagesize=letter,
                                   rightMargin=inch, leftMargin=inch,
                                   topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        DARK   = colors.HexColor("#1A1A2E")
        TIER_COLORS_PDF = [
            colors.HexColor("#2B7FCC"), colors.HexColor("#D4AC0D"),
            colors.HexColor("#7D3C98"), colors.HexColor("#CA6F1E"),
        ]

        title_style = ParagraphStyle("Title2",  parent=styles["Title"],
                                     fontSize=22, textColor=DARK, spaceAfter=6, fontName="Helvetica-Bold")
        h2_style    = ParagraphStyle("H2",      parent=styles["Heading2"],
                                     fontSize=14, textColor=DARK, spaceBefore=18, spaceAfter=6, fontName="Helvetica-Bold")
        h3_style    = ParagraphStyle("H3",      parent=styles["Heading3"],
                                     fontSize=12, spaceBefore=12, spaceAfter=4, fontName="Helvetica-Bold")
        body_style  = ParagraphStyle("Body2",   parent=styles["Normal"],
                                     fontSize=10, leading=16, spaceAfter=6)
        small_style = ParagraphStyle("Small",   parent=styles["Normal"],
                                     fontSize=8, textColor=colors.HexColor("#666666"))

        story = []
        story.append(Paragraph("Supply Chain Intelligence Report", title_style))
        story.append(Paragraph(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}", small_style))
        story.append(Spacer(1, 16))

        story.append(Paragraph("1. Product Overview", h2_style))
        prod_data = [
            ["Product Name",     product.get("product_name", "—")],
            ["Category",         product.get("product_category", "—")],
            ["Industry",         product.get("industry", "—")],
            ["OEM Manufacturer", product.get("oem_manufacturer", "—")],
            ["OEM Country",      product.get("oem_country", "—")],
            ["Key Components",   ", ".join(product.get("key_components", []) or []) or "—"],
        ]
        prod_table = Table(prod_data, colWidths=[2 * inch, 4.5 * inch])
        prod_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F0F0F0")),
            ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#F8F8F8"), colors.white]),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("PADDING",    (0, 0), (-1, -1), 6),
        ]))
        story.append(prod_table)
        story.append(Spacer(1, 12))

        if oems:
            story.append(Paragraph("2. OEM Manufacturers", h2_style))
            oem_header    = ["Company", "Country", "Role", "Market Share", "Notes"]
            oem_rows_data = [oem_header] + [
                [o.get("company_name", ""), o.get("country", ""), o.get("role", ""),
                 o.get("market_share", ""), o.get("notes", "")]
                for o in oems
            ]
            oem_col_w = [1.4 * inch, 0.9 * inch, 1.2 * inch, 0.8 * inch, 2.2 * inch]
            ot = Table(oem_rows_data, colWidths=oem_col_w, repeatRows=1)
            ot.setStyle(TableStyle([
                ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#1A6B3C")),
                ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
                ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",       (0, 0), (-1, -1), 8),
                ("GRID",           (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F0F8F0"), colors.white]),
                ("VALIGN",         (0, 0), (-1, -1), "TOP"),
                ("PADDING",        (0, 0), (-1, -1), 5),
            ]))
            story.append(ot)
            story.append(Spacer(1, 12))

        if summary:
            story.append(Paragraph("3. Executive Summary", h2_style))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 12))

        story.append(Paragraph("4. Supplier Tiers", h2_style))
        tier_names = ["Tier 1", "Tier 2", "Tier 3", "Tier 4"]
        for ti, (tier_key, suppliers) in enumerate(tiers.items()):
            color = TIER_COLORS_PDF[ti] if ti < len(TIER_COLORS_PDF) else colors.grey
            tname = tier_names[ti]       if ti < len(tier_names)      else f"Tier {ti+1}"
            story.append(Paragraph(f"{tname} Suppliers ({len(suppliers)})", h3_style))
            if not suppliers:
                story.append(Paragraph("No suppliers identified.", small_style))
                continue
            header = ["Company", "Country", "Supplies To", "Components", "Confidence"]
            rows   = [header] + [
                [s.get("company_name", "—"), s.get("country", "—"), s.get("supplies_to", "—"),
                 ", ".join((s.get("components_supplied") or [])[:2]), s.get("confidence", "—")]
                for s in suppliers
            ]
            col_w = [1.5 * inch, 1 * inch, 1.5 * inch, 2 * inch, 0.8 * inch]
            t = Table(rows, colWidths=col_w, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",     (0, 0), (-1, 0), color),
                ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
                ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",       (0, 0), (-1, -1), 8),
                ("GRID",           (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F8F8F8"), colors.white]),
                ("VALIGN",         (0, 0), (-1, -1), "TOP"),
                ("PADDING",        (0, 0), (-1, -1), 5),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))

        story.append(Paragraph("5. Statistics", h2_style))
        total     = sum(len(v) for v in tiers.values())
        countries = len(set(s.get("country", "") for arr in tiers.values() for s in arr if s.get("country")))
        stats_data = [
            ["Total Suppliers",  str(total)],
            ["Countries",        str(countries)],
            ["Tier Depth",       str(len(tiers))],
            ["Industry",         product.get("industry", "—")],
        ]
        st = Table(stats_data, colWidths=[3 * inch, 3.5 * inch])
        st.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0), (0, -1), colors.HexColor("#F0F0F0")),
            ("FONTNAME",       (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE",       (0, 0), (-1, -1), 9),
            ("GRID",           (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#F8F8F8"), colors.white]),
            ("PADDING",        (0, 0), (-1, -1), 6),
        ]))
        story.append(st)

        doc.build(story)
        safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", product.get("product_name", "supply_chain"))[:40]
        return Response(
            buffer.getvalue(),
            mimetype="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_supply_chain.pdf"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500
