"""
insights.py — Supply chain insights report generator.

Reads one or more verified Excel files, classifies rows into three tiers,
runs Pandas + NetworkX analysis, generates a Claude narrative, and produces
a downloadable DOCX report.

Row classification:
  Company Exists + Supply Ties + Correct Component = Yes  → fully verified (ground truth)
  Company Exists + Correct Component = Yes, Supply Ties No → model inference (included, flagged ⚠)
  Company Exists + Supply Ties = Yes, Correct Component No → needs research (excluded, listed separately)
  Anything else                                            → excluded entirely

Register in app.py via:
    from insights import insights_bp
    app.register_blueprint(insights_bp)
"""

import io
import json
import os
import re
import tempfile
import threading
import uuid
import datetime
from collections import Counter

import openpyxl
from flask import Blueprint, Response, jsonify, request, send_file

from ai import call_ai, safe_parse_json

insights_bp = Blueprint("insights", __name__)

_insights_queues      = {}
_insights_queues_lock = threading.Lock()
_insights_store       = {}
_insights_store_lock  = threading.Lock()
_insights_file_store  = {}   # key → {"path": str, "filename": str} for per-file docs
_insights_file_lock   = threading.Lock()


# ── SSE helpers ───────────────────────────────────────────────────────────────

def _push(stream_id, event):
    with _insights_queues_lock:
        if stream_id in _insights_queues:
            _insights_queues[stream_id].append(event)

def _status(stream_id, msg):
    _push(stream_id, {"type": "status", "message": msg})

def _log(stream_id, msg, is_error=False):
    _push(stream_id, {"type": "log", "message": msg, "is_error": is_error})


# ── Column detection ──────────────────────────────────────────────────────────

_COL_ALIASES = {
    "company_name":      ["company name", "company"],
    "country":           ["country"],
    "supplies_to":       ["supplies to", "supplies_to"],
    "components":        ["components supplied", "components_supplied"],
    "company_exists":    ["company exists"],
    "supply_ties":       ["supply ties exist", "supply ties"],
    "correct_component": ["correct component supplied", "correct component"],
    "product":           ["product"],
    "url":               ["url", "urls"],
}

_OEM_COL_ALIASES = {
    "company_name":      ["company name", "company"],
    "product":           ["product"],
    "country":           ["country"],
    "regulatory_body":   ["regulatory body"],
    "region":            ["region"],
    "status":            ["status"],
    "details":           ["details"],
    "confidence":        ["confidence"],
}

def _detect_cols(ws, aliases=None):
    """Return {logical_name: 1-based column index} for a worksheet."""
    if aliases is None:
        aliases = _COL_ALIASES
    header = {
        str(cell.value).strip().lower(): cell.column
        for cell in ws[1] if cell.value
    }
    result = {}
    for logical, alias_list in aliases.items():
        for alias in alias_list:
            if alias in header:
                result[logical] = header[alias]
                break
    return result


# ── Data ingestion ────────────────────────────────────────────────────────────

_PRODUCT_OVERVIEW_LABELS = {
    "category", "industry", "oem manufacturer", "oem country",
    "key components", "description",
}

def _read_product_sheet(ws):
    """Read product name, overview fields, and AI summary from a product detail sheet."""
    info = {
        "name": "", "category": "", "industry": "",
        "oem_manufacturer": "", "oem_country": "",
        "key_components": "", "description": "", "ai_summary": "",
    }
    if ws is None:
        return info

    # Row 1 col A = product name (merged cell)
    name_val = ws.cell(1, 1).value
    info["name"] = str(name_val).strip() if name_val else ""

    # Rows 3–8: label (col A) → value (col B)
    capture_next_as_summary = False
    for row in ws.iter_rows(min_row=2, values_only=True):
        label_raw = str(row[0]).strip() if row[0] else ""
        value_raw = str(row[1]).strip() if len(row) > 1 and row[1] else ""

        if capture_next_as_summary:
            # First non-empty merged cell after "Provider:" heading is the AI summary
            # The summary spans col A (col B is blank in a merged cell)
            summary_text = str(row[0]).strip() if row[0] else ""
            if summary_text and not summary_text.startswith("Provider:") \
                    and summary_text.lower() not in ("none", "nan", ""):
                info["ai_summary"] = summary_text
                capture_next_as_summary = False
            continue

        label_lower = label_raw.lower()
        if label_lower in _PRODUCT_OVERVIEW_LABELS:
            key = label_lower.replace(" ", "_")
            info[key] = value_raw
        elif label_lower.startswith("provider:"):
            capture_next_as_summary = True

    return info


def _ingest_files(files_data, stream_id):
    """
    Classify rows from all uploaded files into analysis_rows and research_rows.
    Also reads OEM sheets for regulatory/country data and product sheets for metadata.

    files_data: list of {"filename": str, "bytes": bytes}
    Returns (analysis_rows, research_rows, regulatory_by_file,
             oem_countries_by_file, product_info_by_file).
    """
    analysis_rows        = []
    research_rows        = []
    regulatory_by_file    = {}   # filename → list of regulatory dicts
    oem_countries_by_file = {}   # filename → {oem_name: country}
    product_info_by_file  = {}   # filename → product metadata dict
    file_stats_by_file    = {}   # filename → {tier_num → {total, confirmed, structural, mismatch, inference}}

    for item in files_data:
        filename = item["filename"]
        _status(stream_id, f"📂 Reading {filename}…")

        try:
            wb = openpyxl.load_workbook(io.BytesIO(item["bytes"]), data_only=True)
        except Exception as e:
            _log(stream_id, f"[error] Could not open {filename}: {e}", True)
            continue

        # ── Product detail sheet ─────────────────────────────────────────────
        # The product sheet is the one that is NOT index / oem / tier
        _skip = ("index", "oem", "tier")
        product_ws = next(
            (ws for ws in wb.worksheets
             if not any(kw in ws.title.lower() for kw in _skip)),
            None,
        )
        product_info_by_file[filename] = _read_product_sheet(product_ws)

        # ── Tier sheets ───────────────────────────────────────────────────────
        tier_sheets = [ws for ws in wb.worksheets if "tier" in ws.title.lower()]
        if not tier_sheets:
            _log(stream_id, f"[warn] No tier sheets found in {filename} — skipping", True)
            continue

        for ws in tier_sheets:
            cols = _detect_cols(ws)
            if not cols.get("company_name"):
                _log(stream_id, f"[warn] Sheet '{ws.title}' has no 'Company Name' column — skipping", True)
                continue

            tier_match = re.search(r'tier[_\s]*(\d+)', ws.title.lower())
            tier_num   = int(tier_match.group(1)) if tier_match else 1

            for row in ws.iter_rows(min_row=2, values_only=True):
                def get(key):
                    col = cols.get(key)
                    if col is None:
                        return ""
                    val = row[col - 1]
                    return str(val).strip() if val is not None else ""

                company = get("company_name")
                if not company or company.lower() in ("none", "nan", ""):
                    continue

                exists    = get("company_exists").lower()    == "yes"
                supply    = get("supply_ties").lower()       == "yes"
                component = get("correct_component").lower() == "yes"

                # Track per-tier validation stats for the file summary table
                tier_stats = (file_stats_by_file
                              .setdefault(filename, {})
                              .setdefault(tier_num, {"total": 0, "confirmed": 0,
                                                      "structural": 0, "mismatch": 0,
                                                      "inference": 0}))
                tier_stats["total"] += 1
                if exists and supply and component:
                    tier_stats["confirmed"] += 1
                elif not exists:
                    tier_stats["structural"] += 1
                elif exists and not component:
                    tier_stats["mismatch"] += 1
                elif exists and not supply and component:
                    tier_stats["inference"] += 1

                entry = {
                    "company_name": company,
                    "country":      get("country") or "Unknown",
                    "supplies_to":  get("supplies_to"),
                    "components":   get("components"),
                    "product":      get("product"),
                    "source_file":  filename,
                    "tier":         tier_num,
                    "url":          get("url") or "",
                }

                if exists and supply and component:
                    entry["_inferred"]       = False
                    entry["_needs_research"] = False
                    analysis_rows.append(entry)
                elif exists and component:
                    entry["_inferred"]       = True
                    entry["_needs_research"] = False
                    analysis_rows.append(entry)
                elif exists and supply:
                    entry["_inferred"]       = False
                    entry["_needs_research"] = True
                    research_rows.append(entry)
                # else: excluded entirely

        # ── OEM sheets — country data (all products) + regulatory (pharma) ──────
        oem_sheets      = [ws for ws in wb.worksheets if "oem" in ws.title.lower()]
        reg_rows        = []
        oem_country_map = {}   # company_name → country for this file
        for ws in oem_sheets:
            cols = _detect_cols(ws, _OEM_COL_ALIASES)
            for row in ws.iter_rows(min_row=2, values_only=True):
                def oget(key):
                    col = cols.get(key)
                    if col is None:
                        return ""
                    val = row[col - 1]
                    return str(val).strip() if val is not None else ""
                oem_name = oget("company_name")
                if not oem_name or oem_name.lower() in ("none", "nan", ""):
                    continue
                # Always collect country regardless of pharma status
                country = oget("country")
                if country and country.lower() not in ("unknown", "none", "nan", ""):
                    oem_country_map[oem_name] = country
                # Regulatory data — pharma products only
                reg_body = oget("regulatory_body")
                if reg_body:
                    reg_rows.append({
                        "oem_name":        oem_name,
                        "regulatory_body": reg_body,
                        "region":          oget("region"),
                        "status":          oget("status"),
                        "details":         oget("details"),
                        "confidence":      oget("confidence"),
                        "product_name":    product_info_by_file.get(filename, {}).get("name", ""),
                    })
        if reg_rows:
            regulatory_by_file[filename] = reg_rows
            _log(stream_id, f"[pharma] {filename}: {len(reg_rows)} regulatory row(s) found")
        if oem_country_map:
            oem_countries_by_file[filename] = oem_country_map

    return analysis_rows, research_rows, regulatory_by_file, oem_countries_by_file, product_info_by_file, file_stats_by_file


# ── Pandas analysis ───────────────────────────────────────────────────────────

def _run_pandas(analysis_rows, stream_id):
    import pandas as pd
    _status(stream_id, "📊 Running Pandas analysis…")

    df = pd.DataFrame(analysis_rows)
    total = len(df)

    # Country concentration
    country_counts = (
        df.groupby("country")["company_name"]
        .count()
        .sort_values(ascending=False)
    )
    country_concentration = [
        {"country": c, "count": int(n), "pct": round(n / total * 100, 1)}
        for c, n in country_counts.items()
        if c and c != "Unknown"
    ][:20]

    # Component frequency — split comma-separated values
    all_components = []
    for comp_str in df["components"].dropna():
        all_components.extend(
            c.strip() for c in comp_str.split(",")
            if c.strip() and c.strip().lower() not in ("none", "nan", "")
        )
    comp_counter = Counter(all_components)
    top_components = [
        {"component": c, "count": n}
        for c, n in comp_counter.most_common(15)
    ]

    # Tier distribution
    tier_dist = df.groupby("source_file")["company_name"].count().to_dict()

    # Cross-product overlap (only meaningful if multiple files)
    cross_product = []
    unique_files = df["source_file"].nunique()
    if unique_files > 1:
        for company, grp in df.groupby("company_name"):
            files = grp["source_file"].unique().tolist()
            if len(files) > 1:
                cross_product.append({
                    "company":  company,
                    "products": files,
                    "inferred": bool(grp["_inferred"].any()),
                })
        cross_product.sort(key=lambda x: len(x["products"]), reverse=True)

    return {
        "country_concentration":  country_concentration,
        "top_components":         top_components,
        "tier_distribution":      tier_dist,
        "cross_product_suppliers": cross_product[:10],
    }


# ── NetworkX analysis ─────────────────────────────────────────────────────────

def _run_networkx(analysis_rows, stream_id):
    import networkx as nx
    _status(stream_id, "🕸️ Running graph analysis…")

    G = nx.DiGraph()
    node_inferred = {}
    company_tier  = {}

    for row in analysis_rows:
        src = row["company_name"]
        dst = row["supplies_to"]
        if src and dst:
            G.add_edge(src, dst)
            node_inferred[src] = node_inferred.get(src, False) or row["_inferred"]
            company_tier.setdefault(src, row.get("tier", 1))

    # Hub companies — high out-degree (supplies to many customers)
    out_degree = sorted(G.out_degree(), key=lambda x: x[1], reverse=True)
    hub_companies = [
        {
            "company":           n,
            "supplies_to_count": d,
            "inferred":          node_inferred.get(n, False),
            "tier":              company_tier.get(n, 1),
        }
        for n, d in out_degree[:10] if d > 0
    ]

    # Bottlenecks — high betweenness centrality
    if G.number_of_nodes() > 1:
        bc = nx.betweenness_centrality(G)
        bottleneck_companies = [
            {
                "company":          n,
                "centrality_score": round(s, 4),
                "inferred":         node_inferred.get(n, False),
                "tier":             company_tier.get(n, 1),
            }
            for n, s in sorted(bc.items(), key=lambda x: x[1], reverse=True)[:10]
            if s > 0
        ]
    else:
        bottleneck_companies = []

    # Circular relationships — mutual edges (A supplies B and B supplies A)
    circular_pairs = []
    seen_pairs = set()
    for a, b in G.edges():
        if G.has_edge(b, a):
            key = frozenset([a, b])
            if key not in seen_pairs:
                seen_pairs.add(key)
                circular_pairs.append({
                    "company_a": a,
                    "company_b": b,
                    "inferred":  node_inferred.get(a, False) or node_inferred.get(b, False),
                })

    return {
        "hub_companies":        hub_companies,
        "bottleneck_companies": bottleneck_companies,
        "circular_pairs":       circular_pairs,
    }


# ── Per-product helpers ───────────────────────────────────────────────────────

def _group_by_source(analysis_rows):
    """Group rows by source_file. Returns list of (display_name, rows) sorted by source_file."""
    groups = {}
    for row in analysis_rows:
        groups.setdefault(row["source_file"], []).append(row)
    result = []
    for source_file, rows in sorted(groups.items()):
        names = [r["product"] for r in rows if r.get("product")]
        display = names[0] if names else source_file
        result.append((display, rows))
    return result


def _most_important_oem(rows):
    """Return the OEM: the most common supplies_to among Tier 1 rows.

    Using only Tier 1 rows avoids misidentifying a busy Tier 1 company as the OEM
    when there are many more Tier 2 rows than Tier 1 rows.
    Falls back to global in-degree if no tier information is present.
    """
    tier1_rows = [r for r in rows if r.get("tier") == 1]
    source = tier1_rows if tier1_rows else rows
    counter = Counter(r["supplies_to"] for r in source if r.get("supplies_to"))
    return counter.most_common(1)[0][0] if counter else None


def _analyse_product(rows):
    """Compute per-product hub companies, bottlenecks, and top components."""
    import networkx as nx
    G = nx.DiGraph()
    node_inferred = {}
    company_tier  = {}
    for row in rows:
        src, dst = row["company_name"], row["supplies_to"]
        if src and dst:
            G.add_edge(src, dst)
            node_inferred[src] = node_inferred.get(src, False) or row["_inferred"]
            company_tier.setdefault(src, row.get("tier", 1))

    hub_companies = [
        {"company": n, "supplies_to_count": d, "inferred": node_inferred.get(n, False),
         "tier": company_tier.get(n, 1)}
        for n, d in sorted(G.out_degree(), key=lambda x: x[1], reverse=True)[:5]
        if d > 0
    ]

    if G.number_of_nodes() > 1:
        bc = nx.betweenness_centrality(G)
        bottlenecks = [
            {"company": n, "centrality_score": round(s, 4), "inferred": node_inferred.get(n, False),
             "tier": company_tier.get(n, 1)}
            for n, s in sorted(bc.items(), key=lambda x: x[1], reverse=True)[:5]
            if s > 0
        ]
    else:
        bottlenecks = []

    all_components = []
    for row in rows:
        comp_str = row.get("components", "")
        if comp_str:
            all_components.extend(c.strip() for c in comp_str.split(",") if c.strip())
    top_components = [{"component": c, "count": n} for c, n in Counter(all_components).most_common(10)]

    return {"hub_companies": hub_companies, "bottlenecks": bottlenecks, "top_components": top_components}


def _chart_choropleth(rows, title="Supplier Geographic Concentration", color_scale="Blues"):
    """Choropleth PNG of supplier countries. Pass all rows for combined, or product rows for per-product."""
    try:
        import plotly.express as px
        import pandas as pd
        country_counts = Counter(
            r["country"] for r in rows
            if r.get("country") and r["country"] not in ("Unknown", "")
        )
        if not country_counts:
            return None
        df = pd.DataFrame(list(country_counts.items()), columns=["country", "count"])
        fig = px.choropleth(
            df, locations="country", locationmode="country names",
            color="count", color_continuous_scale=color_scale,
            title=title,
        )
        fig.update_layout(margin={"l": 0, "r": 0, "t": 40, "b": 0}, height=400)
        return fig.to_image(format="png", width=800, height=400)
    except Exception:
        return None


def _chart_sankey(rows, tier_filter=None, title="Supply Chain Flow"):
    """Sankey diagram PNG for one product.

    tier_filter: if given (e.g. 1 or 2), only rows whose 'tier' matches are drawn.
    Pass tier_filter=1 for Tier-1 → OEM, tier_filter=2 for Tier-2 → Tier-1.
    """
    try:
        import plotly.graph_objects as go

        filtered = [r for r in rows if tier_filter is None or r.get("tier") == tier_filter]
        if not filtered:
            return None

        # Collect unique node names preserving insertion order
        node_set, node_index = [], {}
        def _node(name):
            if name not in node_index:
                node_index[name] = len(node_set)
                node_set.append(name)
            return node_index[name]

        tier_colours = {1: "rgba(21,101,192,0.7)", 2: "rgba(230,81,0,0.7)", 3: "rgba(106,27,154,0.7)"}
        link_src, link_tgt, link_val, link_col = [], [], [], []

        for row in filtered:
            src, dst = row.get("company_name", ""), row.get("supplies_to", "")
            if src and dst:
                link_src.append(_node(src))
                link_tgt.append(_node(dst))
                link_val.append(1)
                link_col.append(tier_colours.get(row.get("tier", 1), "rgba(100,100,100,0.4)"))

        if not link_src:
            return None

        # Node colours — destination nodes (OEM for T1 chart, Tier-1 for T2 chart) get a
        # distinct colour so the two columns are visually differentiated
        node_tier = {}
        for row in filtered:
            src = row.get("company_name", "")
            if src:
                node_tier.setdefault(src, row.get("tier", 1))

        dst_nodes = {r.get("supplies_to", "") for r in filtered if r.get("supplies_to")} - \
                    {r.get("company_name", "") for r in filtered if r.get("company_name")}

        dst_colour = "rgba(27,94,32,0.9)"   # green for OEM / Tier-1 destination column
        node_colours = [
            dst_colour if name in dst_nodes
            else tier_colours.get(node_tier.get(name, 1), "rgba(100,100,100,0.7)")
            for name in node_set
        ]

        # Scale layout to node count — keep the chart wide and compact vertically
        n_nodes    = len(node_set)
        fig_h      = max(400, min(1400, n_nodes * 16))  # 16px per node keeps bands slim
        fig_w      = max(1100, min(1800, 1000 + n_nodes * 6))  # wider than tall
        font_size  = max(9, 13 - n_nodes // 12)
        node_pad   = max(3, 10 - n_nodes // 10)
        node_thick = max(8, 12 - n_nodes // 20)  # thinner bars = less "fat"

        _MAX_LBL = 30
        display_labels = [
            n if len(n) <= _MAX_LBL else n[:_MAX_LBL - 1] + "…"
            for n in node_set
        ]

        fig = go.Figure(go.Sankey(
            arrangement="snap",
            node=dict(label=display_labels, color=node_colours,
                      pad=node_pad, thickness=node_thick),
            link=dict(source=link_src, target=link_tgt, value=link_val, color=link_col),
        ))
        fig.update_layout(
            title_text=title,
            font_size=font_size,
            height=fig_h,
            margin={"l": 180, "r": 180, "t": 50, "b": 20},
        )
        return fig.to_image(format="png", width=fig_w, height=fig_h)
    except Exception:
        return None


def _chart_nodemap(rows):
    """Node-link diagram PNG centred on the most important OEM for one product.

    Node tier is computed relative to the OEM, not from the global sheet tier:
      - OEM itself                               → green
      - Directly supplies OEM (relative Tier 1)  → blue
      - Supplies a Tier-1 node (relative Tier 2) → orange
    """
    try:
        import networkx as nx
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import Patch

        oem = _most_important_oem(rows)
        if not oem:
            return None

        rel_tier1 = {r["company_name"] for r in rows if r.get("supplies_to") == oem}
        subgraph_rows = [
            r for r in rows
            if r.get("supplies_to") == oem or r.get("supplies_to") in rel_tier1
        ]
        if not subgraph_rows:
            return None

        G = nx.DiGraph()
        for row in subgraph_rows:
            G.add_edge(row["company_name"], row["supplies_to"])

        def _rel_tier(node):
            if node == oem:       return 0
            if node in rel_tier1: return 1
            return 2

        colour_map = {0: "#1B5E20", 1: "#1565C0", 2: "#E65100"}
        node_list  = list(G.nodes())
        n_nodes    = len(node_list)

        if n_nodes <= 12:
            node_size, font_size, k_spread = 900, 8, 2.0
        elif n_nodes <= 25:
            node_size, font_size, k_spread = 550, 7, 1.8
        else:
            node_size, font_size, k_spread = 320, 6, 1.5

        node_colours = [colour_map.get(_rel_tier(n), "#888888") for n in node_list]
        sizes        = [node_size * 1.6 if n == oem else node_size for n in node_list]

        pos = nx.spring_layout(G, seed=42, k=k_spread)

        fig, ax = plt.subplots(figsize=(10, 7))
        ax.set_facecolor("#F8F8F8")

        nx.draw_networkx_nodes(G, pos, ax=ax, nodelist=node_list,
                               node_color=node_colours, node_size=sizes, alpha=0.9)
        nx.draw_networkx_edges(G, pos, ax=ax,
                               edge_color="#999999", arrows=True, arrowsize=14,
                               width=1.0, connectionstyle="arc3,rad=0.05")
        nx.draw_networkx_labels(G, pos, ax=ax,
                                font_size=font_size, font_color="black", font_weight="bold",
                                bbox=dict(facecolor="white", alpha=0.65,
                                          edgecolor="none", boxstyle="round,pad=0.2"))

        ax.set_title(f"Supply Network — Top OEM: {oem}", fontsize=10, pad=10)
        ax.axis("off")

        legend_elements = [
            Patch(facecolor="#1B5E20", label="OEM"),
            Patch(facecolor="#1565C0", label="Tier 1 (direct supplier to OEM)"),
            Patch(facecolor="#E65100", label="Tier 2 (supplier to Tier 1)"),
        ]
        ax.legend(handles=legend_elements, loc="lower left", fontsize=8,
                  framealpha=0.8, edgecolor="#cccccc")

        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None


def _add_picture_safe(doc, png_bytes, width_inches=5.5):
    """Embed a PNG into the DOCX, silently skip if bytes is None or on error."""
    if not png_bytes:
        return
    try:
        from docx.shared import Inches
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(width_inches))
    except Exception:
        pass


def _generate_leaflet_html(rows, oem_countries=None):
    """Return a self-contained HTML string: Leaflet map with per-company dots and edges.

    All coordinates are resolved in Python via get_coords() + deterministic jitter,
    then embedded as inline JSON — no geocoding happens in the browser.
    Uses CartoDB Positron tiles for a clean, fast-loading background.
    Sets document.title = 'MAP_READY' once the map is fully initialised.
    """
    from coords import get_coords
    import json as _json

    _SKIP = {"", "unknown", "none", "nan", "n/a", "global", "various"}

    def _jitter(name, lat, lng, spread=2.5):
        h = hash(name) & 0xFFFFFF
        dlat = ((h        % 1000) - 500) / 500 * spread
        dlng = (((h >> 8) % 1000) - 500) / 500 * spread
        return round(lat + dlat, 4), round(lng + dlng, 4)

    # ── Build name → jittered (lat, lng) ──────────────────────────────────
    coords_map = {}

    for row in rows:
        name    = (row.get("company_name") or "").strip()
        country = (row.get("country")      or "").strip()
        if not name or country.lower() in _SKIP or name in coords_map:
            continue
        c = get_coords(country)
        if c:
            coords_map[name] = _jitter(name, c[0], c[1])

    for oem_name, country in (oem_countries or {}).items():
        country = (country or "").strip()
        if oem_name not in coords_map and country.lower() not in _SKIP:
            c = get_coords(country)
            if c:
                coords_map[oem_name] = _jitter(oem_name, c[0], c[1])

    if not coords_map:
        return None

    # ── Identify OEM nodes ─────────────────────────────────────────────────
    all_suppliers = {r.get("company_name", "") for r in rows}
    all_customers = {r.get("supplies_to",   "") for r in rows}
    oem_node_names = all_customers - all_suppliers

    # ── Node styling ──────────────────────────────────────────────────────────
    # OEM: red  Tier 1: cyan  Tier 2: dark gold (#B8860B) — readable on grey map
    tier_fill   = {"oem": "#d50c0c", 1: "#47c8ff", 2: "#B8860B"}
    tier_border = {"oem": "#e8ff47", 1: "rgba(0,0,0,0.5)", 2: "rgba(0,0,0,0.6)"}
    tier_weight = {"oem": 2.5,       1: 1.0,                2: 1.2}
    tier_radius = {"oem": 10,        1: 7,                  2: 6}
    # Edge color = same hue as source tier node, semi-transparent
    edge_colors = {"oem": "rgba(213,12,12,0.55)",
                   1:     "rgba(71,200,255,0.60)",
                   2:     "rgba(184,134,11,0.65)"}

    nodes = []
    seen  = set()

    for row in rows:
        name = row.get("company_name", "")
        if name in seen or name not in coords_map:
            continue
        seen.add(name)
        lat, lng = coords_map[name]
        tier = row.get("tier", 1)
        nodes.append({
            "lat": lat, "lng": lng, "name": name,
            "fill":   tier_fill.get(tier, "#888"),
            "border": tier_border.get(tier, "rgba(0,0,0,0.5)"),
            "weight": tier_weight.get(tier, 1.0),
            "radius": tier_radius.get(tier, 6),
        })

    for oem in oem_node_names:
        if oem not in seen and oem in coords_map:
            seen.add(oem)
            lat, lng = coords_map[oem]
            nodes.append({
                "lat": lat, "lng": lng, "name": oem,
                "fill":   tier_fill["oem"],
                "border": tier_border["oem"],
                "weight": tier_weight["oem"],
                "radius": tier_radius["oem"],
            })

    # ── Build edge list ────────────────────────────────────────────────────
    edges = []
    for row in rows:
        src = row.get("company_name", "")
        dst = row.get("supplies_to",   "")
        if src not in coords_map or dst not in coords_map:
            continue
        tier = row.get("tier", 1)
        s_lat, s_lng = coords_map[src]
        d_lat, d_lng = coords_map[dst]
        edges.append({
            "from_lat": s_lat, "from_lng": s_lng,
            "to_lat":   d_lat, "to_lng":   d_lng,
            "color":    edge_colors.get(tier, "rgba(150,150,150,0.45)"),
        })

    if not nodes:
        return None

    nodes_json = _json.dumps(nodes)
    edges_json = _json.dumps(edges)

    html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    /* Greyscale theme: neutral background so coloured nodes/edges pop */
    body {{ background: #e8e8e8; }}
    #map {{ width: 980px; height: 520px; background: #e8e8e8; }}
    .leaflet-container {{ background: #e8e8e8 !important; }}
    /* Apply greyscale to the entire tile pane so ocean + land both lose colour */
    .leaflet-tile-pane {{ filter: grayscale(1) brightness(0.88) contrast(1.05); }}
    .leaflet-control-attribution {{ display: none; }}
  </style>
  <link rel="stylesheet"
        href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
  <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
</head>
<body>
<div id="map">
  <div id="legend" style="
    position:absolute; bottom:18px; right:12px; z-index:1000;
    background:rgba(255,255,255,0.92); border:1px solid #ccc;
    border-radius:5px; padding:10px 14px; font-family:sans-serif;
    font-size:12px; line-height:1.7; box-shadow:0 2px 6px rgba(0,0,0,0.18);
    pointer-events:none;">
    <div style="font-weight:700;font-size:11px;letter-spacing:.06em;
                text-transform:uppercase;color:#444;margin-bottom:6px;">
      Supply Chain
    </div>
    <div style="display:flex;align-items:center;gap:7px;margin-bottom:3px;">
      <span style="display:inline-block;width:13px;height:13px;border-radius:50%;
                   background:#d50c0c;border:2.5px solid #e8ff47;flex-shrink:0;"></span>
      OEM
    </div>
    <div style="display:flex;align-items:center;gap:7px;margin-bottom:3px;">
      <span style="display:inline-block;width:11px;height:11px;border-radius:50%;
                   background:#47c8ff;border:1.5px solid rgba(0,0,0,0.4);flex-shrink:0;"></span>
      Tier 1
    </div>
    <div style="display:flex;align-items:center;gap:7px;margin-bottom:6px;">
      <span style="display:inline-block;width:10px;height:10px;border-radius:50%;
                   background:#B8860B;border:1.5px solid rgba(0,0,0,0.5);flex-shrink:0;"></span>
      Tier 2
    </div>
    <div style="display:flex;align-items:center;gap:7px;">
      <span style="display:inline-block;width:22px;height:0;
                   border-top:2px dashed #888;flex-shrink:0;"></span>
      Supply link
    </div>
  </div>
</div>
<script>
  var nodes = {nodes_json};
  var edges = {edges_json};

  var map = L.map('map', {{ zoomControl: false, attributionControl: false }})
              .setView([20, 10], 2);

  // Same OSM tile source as the frontend
  L.tileLayer('https://{{s}}.tile.openstreetmap.org/{{z}}/{{x}}/{{y}}.png',
    {{ maxZoom: 18 }}).addTo(map);

  // Dashed edges first so nodes render on top
  edges.forEach(function(e) {{
    L.polyline([[e.from_lat, e.from_lng], [e.to_lat, e.to_lng]], {{
      color:     e.color,
      weight:    2,
      opacity:   0.65,
      dashArray: '10 8'
    }}).addTo(map);
  }});

  // Nodes
  var latLngs = [];
  nodes.forEach(function(n) {{
    latLngs.push([n.lat, n.lng]);
    L.circleMarker([n.lat, n.lng], {{
      radius:      n.radius,
      fillColor:   n.fill,
      color:       n.border,
      weight:      n.weight,
      fillOpacity: 0.9,
      opacity:     1
    }}).bindTooltip(n.name).addTo(map);
  }});

  // Fit map to the spread of nodes with padding
  if (latLngs.length > 0) {{
    map.fitBounds(latLngs, {{ padding: [30, 30] }});
  }}

  // Signal readiness after tiles have a moment to load
  setTimeout(function() {{ document.title = 'MAP_READY'; }}, 2500);
</script>
</body>
</html>"""
    return html


def _screenshot_leaflet(html_content, browser=None):
    """Render html_content in a headless Chromium and return a PNG bytes object.

    Pass an open Playwright browser instance to reuse it across multiple calls
    (avoids the ~2s launch cost per product). If browser is None, a temporary
    one is launched and closed inside this call.
    """
    import tempfile, os, pathlib

    tmp_path = None
    own_browser = browser is None
    try:
        from playwright.sync_api import sync_playwright

        # Write HTML to a temp file so Leaflet can load it via file:// URL
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html",
                                         delete=False, encoding="utf-8") as f:
            f.write(html_content)
            tmp_path = f.name

        file_url = pathlib.Path(tmp_path).as_uri()

        def _capture(br):
            page = br.new_page(viewport={"width": 980, "height": 520})
            try:
                page.goto(file_url, wait_until="networkidle", timeout=20_000)
                # Wait until the JS sets the title after the 2.5 s tile delay
                page.wait_for_function(
                    "document.title === 'MAP_READY'", timeout=15_000
                )
                png = page.locator("#map").screenshot()
                return png
            finally:
                page.close()

        if own_browser:
            with sync_playwright() as pw:
                br = pw.chromium.launch(headless=True,
                                        args=["--no-sandbox",
                                              "--disable-dev-shm-usage"])
                try:
                    return _capture(br)
                finally:
                    br.close()
        else:
            return _capture(browser)

    except Exception:
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)




def _add_product_section(doc, helpers, display_name, rows, sub_num,
                         regulatory_rows=None, oem_countries=None, browser=None):
    """Render one product sub-section: metrics tables + optional regulatory + Sankey + node-link."""
    _ct       = helpers["_ct"]
    _hdr_row  = helpers["_hdr_row"]
    _flag_row = helpers["_flag_row"]
    _set_bg   = helpers["_set_bg"]

    doc.add_heading(f"{sub_num}. {display_name}", level=2)

    metrics = _analyse_product(rows)
    top_oem = _most_important_oem(rows)

    # ── 1. Identified top OEM + supply network diagram ────────────────────────
    if top_oem:
        doc.add_heading(f"Identified Top OEM: {top_oem}", level=3)
        doc.add_paragraph(
            f"{top_oem} is the primary OEM for this product, identified as the most common "
            f"destination node among Tier 1 suppliers."
        )
        nodemap_png = _chart_nodemap(rows)
        if nodemap_png:
            doc.add_heading("Supply Network Diagram", level=3)
            _add_picture_safe(doc, nodemap_png, 5.5)
        doc.add_paragraph()

    # ── 2. Tier 1 and Tier 2 supplier tables ─────────────────────────────────
    if top_oem:
        tier1_supplier_rows = [r for r in rows if r.get("supplies_to") == top_oem and r.get("tier") == 1]
        tier1_names         = {r["company_name"] for r in tier1_supplier_rows}
        tier2_supplier_rows = [r for r in rows if r.get("tier") == 2 and r.get("supplies_to") in tier1_names]

        if tier1_supplier_rows:
            doc.add_heading(f"Tier 1 Suppliers to {top_oem}", level=3)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            _hdr_row(tbl, ["Company", "Country", "Components Supplied"], "1565C0")
            for r in tier1_supplier_rows:
                row = tbl.add_row().cells
                label = f"⚠ {r['company_name']}" if r.get("_inferred") else r["company_name"]
                _ct(row[0], label)
                _ct(row[1], r.get("country", ""))
                _ct(row[2], r.get("components", ""))
                if r.get("_inferred"):
                    _set_bg(row[0], "FFFDE7")
            doc.add_paragraph()

        if tier2_supplier_rows:
            doc.add_heading(f"Tier 2 Suppliers (via Tier 1 → {top_oem})", level=3)
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Table Grid"
            _hdr_row(tbl, ["Company", "Country", "Supplies To (Tier 1)", "Components Supplied"], "6A1B9A")
            for r in tier2_supplier_rows:
                row = tbl.add_row().cells
                label = f"⚠ {r['company_name']}" if r.get("_inferred") else r["company_name"]
                _ct(row[0], label)
                _ct(row[1], r.get("country", ""))
                _ct(row[2], r.get("supplies_to", ""))
                _ct(row[3], r.get("components", ""))
                if r.get("_inferred"):
                    _set_bg(row[0], "FFFDE7")
            doc.add_paragraph()

    # ── 3. Hub companies ──────────────────────────────────────────────────────
    if metrics["hub_companies"]:
        doc.add_heading("Hub Companies", level=3)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Company", "Tier", "Customers Supplied", "Note"], "1565C0")
        for entry in metrics["hub_companies"]:
            row = tbl.add_row().cells
            _flag_row(row, entry)
            _ct(row[1], f"Tier {entry.get('tier', '?')}")
            _ct(row[2], entry["supplies_to_count"])
        doc.add_paragraph()

    # ── 4. Bottlenecks ────────────────────────────────────────────────────────
    if metrics["bottlenecks"]:
        doc.add_heading("Potential Bottlenecks", level=3)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Company", "Tier", "Centrality Score", "Note"], "6A1B9A")
        for entry in metrics["bottlenecks"]:
            row = tbl.add_row().cells
            _flag_row(row, entry)
            _ct(row[1], f"Tier {entry.get('tier', '?')}")
            _ct(row[2], entry["centrality_score"])
        doc.add_paragraph()

    # ── 5. Key components ─────────────────────────────────────────────────────
    if metrics["top_components"]:
        doc.add_heading("Key Components", level=3)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Component", "Frequency"], "1B5E20")
        for entry in metrics["top_components"]:
            row = tbl.add_row().cells
            _ct(row[0], entry["component"])
            _ct(row[1], entry["count"])
        doc.add_paragraph()

    # ── 6. Regulatory status (pharma only) ───────────────────────────────────
    if regulatory_rows:
        doc.add_heading("Regulatory Status (OEM Manufacturers)", level=3)

        _STATUS_COLOURS = {
            "approved":   ("C8E6C9", "1B5E20"),
            "registered": ("C8E6C9", "1B5E20"),
            "pending":    ("FFF9C4", "F57F17"),
            "not found":  ("FFCDD2", "B71C1C"),
        }

        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["OEM Manufacturer", "Regulatory Body", "Region", "Status", "Confidence"], "4A148C")

        for entry in regulatory_rows:
            row = tbl.add_row().cells
            _ct(row[0], entry["oem_name"])
            _ct(row[1], entry["regulatory_body"])
            _ct(row[2], entry["region"])
            status_lower = entry["status"].lower() if entry["status"] else ""
            colours = _STATUS_COLOURS.get(status_lower, ("F5F5F5", "000000"))
            _ct(row[3], entry["status"])
            _set_bg(row[3], colours[0])
            _ct(row[4], entry["confidence"])
        doc.add_paragraph()

    # ── 7. Geographic distribution (one choropleth per tier) ─────────────────
    tier1_rows    = [r for r in rows if r.get("tier") == 1]
    tier2_rows    = [r for r in rows if r.get("tier") == 2]
    oem_ctry_rows = [{"country": c} for c in (oem_countries or {}).values()
                     if c and c.lower() not in ("unknown", "none", "nan", "")]

    choropleth_specs = [
        (oem_ctry_rows, "Greens",  f"OEM Manufacturers — {display_name}"),
        (tier1_rows,    "Blues",   f"Tier 1 Suppliers — {display_name}"),
        (tier2_rows,    "Purples", f"Tier 2 Suppliers — {display_name}"),
    ]
    any_choropleth = False
    for chart_rows, scale, chart_label in choropleth_specs:
        if chart_rows:
            png = _chart_choropleth(chart_rows, title=chart_label, color_scale=scale)
            if png:
                if not any_choropleth:
                    doc.add_heading("Geographic Distribution", level=3)
                    any_choropleth = True
                _add_picture_safe(doc, png, 5.5)
    if any_choropleth:
        doc.add_paragraph()

    # ── 8. Supply chain flow (Sankey) ─────────────────────────────────────────
    sankey_png = _chart_sankey(rows)
    if sankey_png:
        doc.add_heading("Supply Chain Flow", level=3)
        _add_picture_safe(doc, sankey_png, 5.5)
        doc.add_paragraph()

    # ── 9. Supplier world map (Leaflet screenshot via headless Chromium) ─────────
    leaflet_html = _generate_leaflet_html(rows, oem_countries=oem_countries)
    if leaflet_html:
        geo_png = _screenshot_leaflet(leaflet_html, browser=browser)
        if geo_png:
            doc.add_heading("Supplier World Map", level=3)
            _add_picture_safe(doc, geo_png, 5.5)
            doc.add_paragraph()


# ── Claude narrative ──────────────────────────────────────────────────────────

def _generate_narrative(metrics, stream_id, regulatory_by_file=None):
    _status(stream_id, "🤖 Generating narrative with Claude…")

    prompt = f"""You are a senior supply chain intelligence analyst writing a formal briefing document.
Below is a structured JSON object containing quantitative metrics computed from one or more
verified supplier datasets. Write substantive, detailed analysis for each section.

Rules:
- Every claim MUST cite a specific number, percentage, or company name from the metrics JSON.
  Do not introduce facts or companies not present in the data.
- Each section should be 2-3 full paragraphs of analytical prose.
- Within each section, cite concrete examples from the data (e.g. specific company names,
  specific countries, specific components, specific centrality scores).
- Where relevant, explain the strategic implication of what the numbers mean — not just
  what they are, but why they matter for supply chain resilience.
- Use plain prose paragraphs separated by \\n\\n. Do not use bullet points or headers inside
  the text values.

METRICS:
{json.dumps(metrics, indent=2)}

Return ONLY valid JSON (no markdown, no code fences) with these exact keys:
{{
  "executive_summary": "2-3 paragraphs: overall picture of the supply chain dataset — total verified vs inferred suppliers, which products were analysed, the dominant geography, the most critical component category, and the single most important risk finding. Name specific companies and countries.",
  "geo_concentration": "2-3 paragraphs: deep analysis of geographic concentration. Which countries dominate and by what percentage? What does that imply for geopolitical risk (sanctions, trade restrictions, natural disasters)? Cite the top 3-5 countries by name and count. Discuss what a disruption in the top country would mean for the overall supply chain.",
  "key_components": "2-3 paragraphs: which components appear most frequently across the supply chain and what does that tell us? Are there components supplied by only a small number of companies (concentration risk)? Cite the top components by name and frequency. Discuss substitutability and criticality.",
  "hub_companies": "2-3 paragraphs: identify the most connected suppliers by out-degree. Name the top hub companies and their customer counts. Explain why a hub company going offline would cascade across the supply chain. Note which of these are model-inferred relationships vs fully verified.",
  "bottlenecks": "2-3 paragraphs: analyse the betweenness centrality results. Name the top bottleneck companies and their centrality scores. Explain in plain terms what betweenness centrality means — a company with high betweenness sits on many of the shortest paths through the supply network, so its failure disconnects large portions of the chain. Recommend mitigation strategies (dual sourcing, safety stock, etc).",
  "circular_relationships": "2-3 paragraphs: analyse any circular supply relationships detected (Company A supplies Company B AND Company B supplies Company A). Name the specific pairs. Explain why circular relationships are anomalous in a supply chain — they may indicate data quality issues, legitimate cross-supply agreements, or joint venture structures that require scrutiny. Discuss the operational and financial risk these create (circular dependency, price-fixing exposure, single-point-of-failure amplification). Write empty string if no circular pairs were detected.",
  "cross_product": "2-3 paragraphs on suppliers shared across multiple products, naming specific companies and the products they serve. Explain the dual risk: shared suppliers increase efficiency but create correlated failure — a disruption hits multiple product lines at once. Write empty string if only 1 product or file was analysed.",
  "needs_research": "2 paragraphs: summarise the entries that require further research (Company Exists + Supply Ties confirmed but component unverified). Name specific companies if present. Recommend concrete next steps — targeted web searches, direct outreach, or procurement team verification. Write empty string if the list is empty."
}}"""

    # Collect pharma regulatory rows (used by third call below)
    all_reg_rows = []
    if regulatory_by_file:
        for rows in regulatory_by_file.values():
            all_reg_rows.extend(rows)

    metrics_json = json.dumps(metrics, indent=2)

    prompt_a = f"""You are a senior supply chain intelligence analyst writing a formal briefing document.
Below is a structured JSON object containing quantitative metrics from one or more verified supplier datasets.
Write substantive, detailed analytical prose for each section.

Rules:
- Every claim MUST cite a specific number, percentage, or company name from the metrics JSON.
- Each section should be 2-3 full paragraphs of analytical prose separated by \\n\\n.
- Do not use bullet points or headers inside the text values.

METRICS:
{metrics_json}

Return ONLY valid JSON (no markdown, no code fences) with these exact keys:
{{
  "executive_summary": "2-3 paragraphs: overall picture — total verified vs inferred suppliers, products analysed, dominant geography, most critical component category, and the single most important risk finding. Name specific companies and countries.",
  "geo_concentration": "2-3 paragraphs: which countries dominate and by what percentage? Geopolitical risk implications (sanctions, trade restrictions). Cite top 3-5 countries by name and count. What would a disruption in the top country mean?",
  "key_components": "2-3 paragraphs: which components appear most frequently? Are there components supplied by very few companies? Cite top components by name and frequency. Discuss substitutability and criticality.",
  "hub_companies": "2-3 paragraphs: identify the most connected suppliers by out-degree. Name top hub companies and their customer counts. Explain cascade risk. Note which relationships are inferred vs verified."
}}"""

    prompt_b = f"""You are a senior supply chain intelligence analyst writing a formal briefing document.
Below is a structured JSON object containing quantitative metrics from one or more verified supplier datasets.
Write substantive, detailed analytical prose for each section.

Rules:
- Every claim MUST cite a specific number, percentage, or company name from the metrics JSON.
- Each section should be 2-3 full paragraphs of analytical prose separated by \\n\\n.
- Do not use bullet points or headers inside the text values.

METRICS:
{metrics_json}

Return ONLY valid JSON (no markdown, no code fences) with these exact keys:
{{
  "bottlenecks": "2-3 paragraphs: analyse betweenness centrality results. Name top bottleneck companies and their scores. Explain in plain terms what betweenness centrality means. Recommend mitigation strategies.",
  "circular_relationships": "2-3 paragraphs: analyse any circular supply relationships detected. Name specific pairs. Explain why circularity is anomalous and the risks it creates. Write empty string if none detected.",
  "cross_product": "2-3 paragraphs: suppliers shared across multiple products — name specific companies and products. Explain the dual risk of shared suppliers (efficiency vs correlated failure). Write empty string if only 1 product analysed.",
  "needs_research": "2 paragraphs: summarise entries requiring further research. Name specific companies. Recommend concrete next steps. Write empty string if none."
}}"""

    _status(stream_id, "🤖 Generating narrative part 1/2…")
    raw_a  = call_ai(prompt_a, "anthropic", max_tokens=6000)
    part_a = safe_parse_json(raw_a)

    _status(stream_id, "🤖 Generating narrative part 2/2…")
    raw_b  = call_ai(prompt_b, "anthropic", max_tokens=6000)
    part_b = safe_parse_json(raw_b)

    # Merge — fall back to empty string for any section that failed to parse
    _EMPTY = {"executive_summary": "", "geo_concentration": "", "key_components": "",
              "hub_companies": "", "bottlenecks": "", "circular_relationships": "",
              "cross_product": "", "needs_research": "", "regulatory_analysis": ""}

    result = {**_EMPTY}
    if isinstance(part_a, dict):
        result.update({k: v for k, v in part_a.items() if k in _EMPTY})
    if isinstance(part_b, dict):
        result.update({k: v for k, v in part_b.items() if k in _EMPTY})

    # If part_a failed entirely, surface the raw text in executive_summary so nothing is silently lost
    if not isinstance(part_a, dict) or "executive_summary" not in part_a:
        result["executive_summary"] = raw_a if not isinstance(part_a, dict) else result["executive_summary"]

    # Pharma regulatory section — third call if needed
    if all_reg_rows:
        _status(stream_id, "🤖 Generating regulatory narrative…")
        reg_json_str = json.dumps(all_reg_rows, indent=2)
        prompt_reg = f"""You are a senior supply chain intelligence analyst.
Analyse the regulatory approval landscape for the OEM manufacturers in this pharmaceutical supply chain.

REGULATORY DATA:
{reg_json_str[:3000]}

Return ONLY valid JSON (no markdown):
{{
  "regulatory_analysis": "2-3 paragraphs: which regulatory bodies are represented (FDA, EMA, PMDA, etc.)? Which OEMs are fully approved vs pending vs not found? What does the approval pattern imply for market access, compliance risk, and supply chain resilience? Cite specific OEM names, regulatory bodies, and status values."
}}"""
        raw_reg  = call_ai(prompt_reg, "anthropic", max_tokens=2000)
        part_reg = safe_parse_json(raw_reg)
        if isinstance(part_reg, dict) and "regulatory_analysis" in part_reg:
            result["regulatory_analysis"] = part_reg["regulatory_analysis"]

    return result


# ── Per-file verified supply data report ─────────────────────────────────────

def _build_per_file_docx(product_info, file_rows, research_rows_for_file,
                         reg_rows=None, oem_countries=None):
    """Build a standalone Word doc for one uploaded Excel file.

    Includes product overview, AI summary, tier tables with a
    Verified / Model Inference status column, and regulatory landscape if present.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_bg(cell, hex6):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex6)
        tcPr.append(shd)

    def _ct(cell, text, bold=False, size=10, hex_color=None):
        para = cell.paragraphs[0]
        run  = para.add_run(str(text) if text is not None else "—")
        run.bold      = bold
        run.font.size = Pt(size)
        if hex_color:
            b = bytes.fromhex(hex_color)
            run.font.color.rgb = RGBColor(b[0], b[1], b[2])

    def _hdr_row(tbl, headers, bg_hex):
        r = tbl.rows[0].cells
        for i, h in enumerate(headers):
            _set_bg(r[i], bg_hex)
            _ct(r[i], h, bold=True, hex_color="FFFFFF")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    # Title
    product_name = product_info.get("name") or "Unknown Product"
    t = doc.add_heading(f"{product_name} — Verified Supply Data Report", 0)
    t.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    date_para = doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M:%S')}")
    date_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    verified_count = sum(1 for r in file_rows if not r.get("_inferred"))
    inferred_count = sum(1 for r in file_rows if r.get("_inferred"))
    tier1_count    = sum(1 for r in file_rows if r.get("tier") == 1)
    tier2_count    = sum(1 for r in file_rows if r.get("tier") == 2)
    oem_names      = {r["supplies_to"] for r in file_rows if r.get("tier") == 1 and r.get("supplies_to")}
    oem_count      = len(oem_names)

    counts_data = [
        ("OEM Manufacturers", oem_count),
        ("Tier 1 Suppliers",  tier1_count),
        ("Tier 2 Suppliers",  tier2_count),
        ("Fully Verified",    verified_count),
        ("Model Inference",   inferred_count),
    ]
    counts_tbl = doc.add_table(rows=len(counts_data), cols=2)
    counts_tbl.style = "Table Grid"
    for i, (label, val) in enumerate(counts_data):
        cells = counts_tbl.rows[i].cells
        _ct(cells[0], label, bold=True)
        _set_bg(cells[0], "E8EAF6")
        _ct(cells[1], val)
    doc.add_paragraph()

    # Product overview
    overview = [
        ("Category",         product_info.get("category", "")),
        ("Industry",         product_info.get("industry", "")),
        ("OEM Manufacturer", product_info.get("oem_manufacturer", "")),
        ("OEM Country",      product_info.get("oem_country", "")),
        ("Key Components",   product_info.get("key_components", "")),
        ("Description",      product_info.get("description", "")),
    ]
    has_overview = any(v for _, v in overview)
    if has_overview:
        doc.add_heading("1. Product Overview", level=1)
        tbl = doc.add_table(rows=len(overview), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(overview):
            cells = tbl.rows[i].cells
            _ct(cells[0], label, bold=True)
            _set_bg(cells[0], "E8EAF6")
            _ct(cells[1], value)
        doc.add_paragraph()

    # AI summary
    ai_summary = product_info.get("ai_summary", "")
    if ai_summary:
        doc.add_heading("2. AI Analysis Summary", level=1)
        doc.add_paragraph(ai_summary)
        doc.add_paragraph()

    next_sec = 3

    # ── OEM Manufacturers table ───────────────────────────────────────────────
    oem_names_ordered = []
    seen_oems = set()
    for r in file_rows:
        oem = r.get("supplies_to", "")
        if r.get("tier") == 1 and oem and oem not in seen_oems:
            seen_oems.add(oem)
            oem_names_ordered.append(oem)

    if oem_names_ordered:
        doc.add_heading(f"{next_sec}. OEM Manufacturers", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["OEM Manufacturer", "Country"], "1A237E")
        _oem_ctry = oem_countries or {}
        for oem in oem_names_ordered:
            cells = tbl.add_row().cells
            _ct(cells[0], oem, bold=True)
            _ct(cells[1], _oem_ctry.get(oem, ""))
        doc.add_paragraph()
        next_sec += 1

    # ── Tier tables ───────────────────────────────────────────────────────────
    _TIER_COLORS = {1: "1565C0", 2: "6A1B9A"}

    def _set_col_widths(tbl, widths_inches):
        """Force exact column widths via XML on every cell in the table.

        Must be called AFTER all rows have been added — it iterates over
        every row so new rows pick up the correct widths too.
        Also disables Word's autofit so the widths are not overridden on open.
        """
        tbl.allow_autofit = False
        for row in tbl.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx >= len(widths_inches):
                    break
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"),    str(int(widths_inches[col_idx] * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

    def _add_url_lines(cell, raw_url_str, size=7.5):
        """Write the first 4 URLs (company-exists results) one per line.
        Inserts zero-width spaces after '/', '?' and '&' so Word can line-break
        long URLs at sensible points without the text overflowing the column.
        """
        if not raw_url_str:
            _ct(cell, "—", size=size)
            return
        urls = [u.strip() for u in raw_url_str.split(",") if u.strip()][:4]
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(2)
        for i, url in enumerate(urls):
            if i > 0:
                para.add_run("\n")
            # Insert zero-width space (​) after break-friendly characters
            breakable = url.replace("/", "/​").replace("?", "?​").replace("&", "&​").replace("=", "=​").replace("-", "-​")
            run = para.add_run(breakable)
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    # Column widths (inches) — total 6.5" for letter page with 1" margins each side
    # Company | Country | Supplies To | Components | Status | URLs
    _TIER_COL_WIDTHS = [1.2, 0.75, 1.0, 1.05, 0.65, 1.85]

    for tier_num in (1, 2):
        tier_rows = [r for r in file_rows if r.get("tier") == tier_num]
        if not tier_rows:
            continue
        doc.add_heading(f"{next_sec}. Tier {tier_num} Suppliers", level=1)
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        _hdr_row(tbl,
                 ["Company", "Country", "Supplies To", "Components Supplied",
                  "Status", "Verification URLs"],
                 _TIER_COLORS.get(tier_num, "37474F"))
        for row in tier_rows:
            cells = tbl.add_row().cells
            is_inferred = row.get("_inferred", False)
            label_name  = f"⚠ {row['company_name']}" if is_inferred else row["company_name"]
            _ct(cells[0], label_name, size=9)
            _ct(cells[1], row.get("country", ""), size=9)
            _ct(cells[2], row.get("supplies_to", ""), size=9)
            _ct(cells[3], row.get("components", ""), size=9)
            status = "Model Inference" if is_inferred else "Verified"
            _ct(cells[4], status, size=9)
            if is_inferred:
                _set_bg(cells[4], "FFFDE7")
            else:
                _set_bg(cells[4], "C8E6C9")
            _add_url_lines(cells[5], row.get("url", ""))
        # Apply widths after all rows exist so every row cell gets set
        _set_col_widths(tbl, _TIER_COL_WIDTHS)
        doc.add_paragraph()
        next_sec += 1

    # Regulatory landscape (pharma only — present when OEM sheet has regulatory_body column)
    if reg_rows:
        doc.add_heading(f"{next_sec}. Regulatory Landscape", level=1)
        _REG_STATUS_COLOURS = {
            "approved":   "C8E6C9",
            "registered": "C8E6C9",
            "pending":    "FFF9C4",
            "not found":  "FFCDD2",
        }
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["OEM Manufacturer", "Regulatory Body", "Region", "Status", "Confidence"], "4A148C")
        for entry in reg_rows:
            cells = tbl.add_row().cells
            _ct(cells[0], entry.get("oem_name", ""))
            _ct(cells[1], entry.get("regulatory_body", ""))
            _ct(cells[2], entry.get("region", ""))
            _ct(cells[3], entry.get("status", ""))
            bg = _REG_STATUS_COLOURS.get((entry.get("status") or "").lower(), "F5F5F5")
            _set_bg(cells[3], bg)
            _ct(cells[4], entry.get("confidence", ""))
        doc.add_paragraph()
        next_sec += 1

    # Requires further research
    if research_rows_for_file:
        doc.add_heading(f"{next_sec}. Requires Further Research", level=1)
        doc.add_paragraph(
            "The following entries have confirmed company existence and supply ties "
            "but the component supplied could not be verified."
        )
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Company", "Supplies To", "Country"], "BF360C")
        for row in research_rows_for_file:
            cells = tbl.add_row().cells
            _ct(cells[0], row["company_name"])
            _ct(cells[1], row.get("supplies_to", ""))
            _ct(cells[2], row.get("country", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _build_docx(metrics, narrative, analysis_rows, research_rows,  # noqa: C901
                regulatory_by_file=None, oem_countries_by_file=None,
                file_stats_by_file=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    def _set_bg(cell, hex6):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex6)
        tcPr.append(shd)

    def _ct(cell, text, bold=False, size=10, hex_color=None):
        para = cell.paragraphs[0]
        run  = para.add_run(str(text) if text is not None else "—")
        run.bold      = bold
        run.font.size = Pt(size)
        if hex_color:
            b = bytes.fromhex(hex_color)
            run.font.color.rgb = RGBColor(b[0], b[1], b[2])

    def _hdr_row(tbl, headers, bg_hex):
        r = tbl.rows[0].cells
        for i, h in enumerate(headers):
            _set_bg(r[i], bg_hex)
            _ct(r[i], h, bold=True, hex_color="FFFFFF")

    def _flag_row(cells, entry, name_key="company"):
        name  = entry.get(name_key, "")
        label = f"⚠ {name}" if entry.get("inferred") else name
        _ct(cells[0], label)
        if entry.get("inferred"):
            _ct(cells[-1], "Model inference")
            _set_bg(cells[-1], "FFFDE7")

    def _add_narrative(text: str):
        """Split on double-newlines and add each chunk as a separate DOCX paragraph."""
        if not text:
            return
        for chunk in str(text).split("\n\n"):
            chunk = chunk.strip()
            if chunk:
                doc.add_paragraph(chunk)

    # Pass closures to per-product section renderer
    _docx_helpers = {"_ct": _ct, "_hdr_row": _hdr_row, "_flag_row": _flag_row, "_set_bg": _set_bg}

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    # ── Title ────────────────────────────────────────────────────────────────
    t = doc.add_heading("Supply Chain Insights Report", 0)
    t.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    date_para = doc.add_paragraph(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
    date_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    for line in [
        f"Files analysed: {', '.join(metrics['products_analysed'])}",
        f"Fully verified suppliers: {metrics['total_verified_suppliers']}",
        f"Model inference suppliers: {metrics['total_inferred_suppliers']}",
        f"Entries requiring further research: {len(research_rows)}",
    ]:
        doc.add_paragraph(line).runs[0].font.size = Pt(10)

    # ── Executive Summary ────────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    _add_narrative(narrative.get("executive_summary", ""))

    # ── Validation Methodology & Findings ────────────────────────────────────
    doc.add_heading("2. Validation Methodology & Findings", level=1)
    doc.add_paragraph(
        "Each entry generated by the discovery pipeline was independently verified through an "
        "automated web-crawl and LLM-judge process. Verification assessed three questions for "
        "each supplier row and produced a binary Yes/No verdict for each."
    )

    doc.add_heading("Data Generation Process", level=2)
    doc.add_paragraph(
        "Supply chain data was generated using an agentic pipeline (Gemini 2.5 Flash for "
        "non-China suppliers, DeepSeek Chat for China-based suppliers) that identified OEM "
        "manufacturers, Tier 1 direct suppliers, and Tier 2 sub-suppliers for each product. "
        "Both models ran in parallel at every step and their outputs were merged and deduplicated "
        "before moving on. Prior to each LLM call, live web evidence was retrieved via DuckDuckGo "
        "and scraped with Crawl4AI (up to 4 pages, 2 000 characters each) and injected into the "
        "prompt — grounding outputs in current web data rather than solely training knowledge. "
        "All LLM outputs were treated as unverified hypotheses and passed to the verification "
        "pipeline before use in any analysis."
    )
    doc.add_paragraph(
        "Each entry in the exported workbook contains: Company Name, Country, Supplies To, "
        "OEM Root, Components Supplied, Confidence (high/medium/low), and Source Hint."
    )

    doc.add_heading("Automated Verification Pipeline", level=2)
    doc.add_paragraph(
        "The verification pipeline answered three independent questions for each supplier row "
        "using web search evidence. Three separate DuckDuckGo searches were run, each scraping "
        "up to 4 pages via Crawl4AI. The combined evidence (up to ~6 000 characters) was sent "
        "to an LLM in a single call that answered all three questions simultaneously. Each "
        "verdict was labelled with its evidence source — web_evidence if a confirming source "
        "was found, or training_knowledge if the model fell back to its parametric knowledge. "
        "All rows across all tier sheets were processed in parallel."
    )
    vp_tbl = doc.add_table(rows=1, cols=3)
    vp_tbl.style = "Table Grid"
    _hdr_row(vp_tbl, ["Question", "Search Query Used", "Output Field"], "1A237E")
    _vp_rows = [
        ('Q1 — Does this company exist?',
         '"{company_name}" supplier manufacturer',
         'Company Exists = Yes / No'),
        ('Q2 — Is there a supply relationship between this company and its stated customer?',
         '"{company_name}" "{supplies_to}" supply partnership',
         'Supply Ties = Yes / No'),
        ('Q3 — Does this company supply the stated component?',
         '"{company_name}" "{components_supplied}" manufacture supply',
         'Correct Component = Yes / No'),
    ]
    for q, query, output in _vp_rows:
        row = vp_tbl.add_row().cells
        _ct(row[0], q)
        _ct(row[1], query)
        _ct(row[2], output)
    doc.add_paragraph()

    doc.add_heading("Entry Classification Framework", level=2)
    doc.add_paragraph(
        "Each row was classified into one of three categories based on the combination of "
        "verdicts across the three verification questions:"
    )
    cls_tbl = doc.add_table(rows=1, cols=3)
    cls_tbl.style = "Table Grid"
    _hdr_row(cls_tbl, ["Classification", "Criteria", "Treatment in This Report"], "1A237E")
    _clf_rows = [
        ("Confirmed Entry (3 Yes)",
         "Company Exists = Yes, Supply Ties = Yes, Correct Component = Yes. "
         "Web evidence found for all three questions.",
         "Used as ground truth in all calculations.",
         "C8E6C9"),
        ("Strong Entry (CE + CC)",
         "Company Exists = Yes, Correct Component = Yes. "
         "Supply Ties could not be corroborated by a public web source. "
         "The supplier is real and produces the correct component, but the specific commercial "
         "relationship is undocumented — consistent with supply agreements being treated as "
         "commercially sensitive trade secrets.",
         "Included in analysis, flagged ⚠ as model inference.",
         "FFFDE7"),
        ("Excluded / Flagged Entry",
         "Company Exists = No, or Correct Component = No, or the row is structurally invalid.",
         "Excluded from all quantitative analysis.",
         "FFCDD2"),
    ]
    for cls_label, criteria, treatment, bg in _clf_rows:
        row = cls_tbl.add_row().cells
        _ct(row[0], cls_label, bold=True)
        _ct(row[1], criteria)
        _ct(row[2], treatment)
        _set_bg(row[0], bg)
    doc.add_paragraph(
        "The absence of a confirming URL for Supply Ties is not treated as evidence of an "
        "error in the generated data. Specific commercial procurement relationships between "
        "named companies are rarely disclosed publicly — they surface only through press "
        "releases announcing strategic partnerships, joint ventures, equity investments, or "
        "supplier-award announcements. The verification pipeline's inability to find such a "
        "source reflects the fundamental opacity of procurement relationships, not a flaw in "
        "the underlying supply chain entry."
    )
    doc.add_paragraph()

    doc.add_heading("Structural Flags", level=2)
    doc.add_paragraph(
        "In addition to the three-tier quality classification, entries were reviewed for "
        "structural validity. Rows were flagged under the following conditions:"
    )
    for flag_text in [
        "OEM-as-Tier-1 Supplier — a company designated as an OEM end-product manufacturer was listed as a Tier 1 component supplier to another OEM (role confusion in the generated data).",
        "Tier-1-as-Tier-2 Supplier — a company already present in the Tier 1 sheet was duplicated as a Tier 2 supplier (inflates apparent supply chain depth).",
        "Self-Loop — the Supplies To field contains the same entity as the Company Name (structurally invalid relationship).",
        "Tier Misclassification — the described relationship is inverted (a downstream customer listed as supplying to its own vendor).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(flag_text)
    doc.add_paragraph(
        "Flagged rows are excluded from all classification counts and surfaced separately "
        "in the 'Requires Further Research' section."
    )
    doc.add_paragraph()

    # ── Aggregate validation findings table ──────────────────────────────────
    doc.add_heading("Aggregate Validation Findings", level=2)
    doc.add_paragraph(
        "The table below consolidates validation results across all uploaded files, broken down "
        "by tier level. Confirmed = all three validation fields Yes. "
        "Strong = Company Exists and Correct Component both Yes, Supply Ties not confirmed. "
        "Flagged = structural error or component mismatch (excluded from Confirmed/Strong counts). "
        "Percentages are relative to total entries for that tier."
    )

    # Aggregate file_stats_by_file across all files, keyed by tier_num
    agg = {}
    for tier_map in (file_stats_by_file or {}).values():
        for tier_num, ts in tier_map.items():
            a = agg.setdefault(tier_num, {"total": 0, "confirmed": 0, "inference": 0,
                                           "structural": 0, "mismatch": 0})
            for k in a:
                a[k] += ts[k]

    agg_tbl = doc.add_table(rows=1, cols=8)
    agg_tbl.style = "Table Grid"
    _hdr_row(agg_tbl, ["Tier Level", "Total Entries",
                        "Confirmed (3 Yes)", "%",
                        "Strong (CE+CC)", "%",
                        "Flagged", "%"], "1A237E")

    totals = {"total": 0, "confirmed": 0, "inference": 0, "flagged": 0}
    for tier_num in sorted(agg.keys()):
        a    = agg[tier_num]
        tot  = a["total"]
        conf = a["confirmed"]
        inf  = a["inference"]
        flag = a["structural"] + a["mismatch"]
        pct  = lambda n: f"{n/tot*100:.1f}%" if tot else "—"
        row  = agg_tbl.add_row().cells
        _ct(row[0], f"Tier {tier_num}")
        _ct(row[1], str(tot))
        _ct(row[2], str(conf));  _set_bg(row[2], "C8E6C9")
        _ct(row[3], pct(conf))
        _ct(row[4], str(inf));   _set_bg(row[4], "FFFDE7")
        _ct(row[5], pct(inf))
        _ct(row[6], str(flag));  _set_bg(row[6], "FFCDD2")
        _ct(row[7], pct(flag))
        totals["total"]    += tot
        totals["confirmed"] += conf
        totals["inference"] += inf
        totals["flagged"]   += flag

    # TOTAL row
    tt   = totals["total"]
    tpct = lambda n: f"{n/tt*100:.1f}%" if tt else "—"
    row  = agg_tbl.add_row().cells
    _ct(row[0], "TOTAL", bold=True)
    _ct(row[1], str(tt), bold=True)
    _ct(row[2], str(totals["confirmed"]), bold=True);  _set_bg(row[2], "C8E6C9")
    _ct(row[3], tpct(totals["confirmed"]), bold=True)
    _ct(row[4], str(totals["inference"]),  bold=True);  _set_bg(row[4], "FFFDE7")
    _ct(row[5], tpct(totals["inference"]),  bold=True)
    _ct(row[6], str(totals["flagged"]),    bold=True);  _set_bg(row[6], "FFCDD2")
    _ct(row[7], tpct(totals["flagged"]),    bold=True)
    for cell in row:
        _set_bg(cell, "ECEFF1") if cell.text not in ("", "—") else None
    doc.add_paragraph()

    # ── Geographical Concentration ───────────────────────────────────────────
    doc.add_heading("3. Geographical Concentration", level=1)
    _add_narrative(narrative.get("geo_concentration", ""))

    if metrics["country_concentration"]:
        top = metrics["country_concentration"][:10]
        labels = [c["country"] for c in reversed(top)]
        values = [c["count"]   for c in reversed(top)]
        fig, ax = plt.subplots(figsize=(6, max(3, len(labels) * 0.45)))
        ax.barh(labels, values, color="#1A237E")
        ax.set_xlabel("Supplier Count", fontsize=8)
        ax.set_title("Geographical Concentration (Top 10)", fontsize=9)
        ax.tick_params(labelsize=7)
        chart_buf = io.BytesIO()
        fig.savefig(chart_buf, format="png", bbox_inches="tight", dpi=150)
        plt.close(fig)
        chart_buf.seek(0)
        doc.add_picture(chart_buf, width=Inches(5.5))

    choropleth_png = _chart_choropleth(analysis_rows, title="Supplier Geographic Concentration (All Products)")
    _add_picture_safe(doc, choropleth_png, 5.5)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Country", "Supplier Count", "% of Total"], "1A237E")
    for entry in metrics["country_concentration"]:
        row = tbl.add_row().cells
        _ct(row[0], entry["country"])
        _ct(row[1], entry["count"])
        _ct(row[2], f"{entry['pct']}%")
    doc.add_paragraph()

    # ── Key Components ───────────────────────────────────────────────────────
    doc.add_heading("4. Key Components", level=1)
    _add_narrative(narrative.get("key_components", ""))
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Component", "Frequency"], "1B5E20")
    for entry in metrics["top_components"]:
        row = tbl.add_row().cells
        _ct(row[0], entry["component"])
        _ct(row[1], entry["count"])
    doc.add_paragraph()


    # ── Hub Companies ────────────────────────────────────────────────────────
    doc.add_heading("5. Hub Companies", level=1)
    _add_narrative(narrative.get("hub_companies", ""))
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Company", "Tier", "Customers Supplied", "Note"], "1565C0")
    for entry in metrics["hub_companies"]:
        row = tbl.add_row().cells
        _flag_row(row, entry)
        _ct(row[1], f"Tier {entry.get('tier', '?')}")
        _ct(row[2], entry["supplies_to_count"])
    doc.add_paragraph()

    # ── Bottlenecks ──────────────────────────────────────────────────────────
    doc.add_heading("6. Potential Bottlenecks", level=1)
    _add_narrative(narrative.get("bottlenecks", ""))
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Company", "Tier", "Centrality Score", "Note"], "6A1B9A")
    for entry in metrics["bottleneck_companies"]:
        row = tbl.add_row().cells
        _flag_row(row, entry)
        _ct(row[1], f"Tier {entry.get('tier', '?')}")
        _ct(row[2], entry["centrality_score"])
    doc.add_paragraph()

    # ── Regulatory Analysis (pharma only) ────────────────────────────────────
    next_sec = 7
    all_reg_rows = []
    if regulatory_by_file:
        for rows in regulatory_by_file.values():
            all_reg_rows.extend(rows)

    if all_reg_rows:
        doc.add_heading(f"{next_sec}. Regulatory Landscape", level=1)
        reg_narrative = narrative.get("regulatory_analysis", "")
        if reg_narrative:
            _add_narrative(reg_narrative)

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Drug / Product", "OEM Manufacturer", "Regulatory Body", "Region", "Status", "Confidence"], "4A148C")

        _REG_STATUS_COLOURS = {
            "approved":   "C8E6C9",
            "registered": "C8E6C9",
            "pending":    "FFF9C4",
            "not found":  "FFCDD2",
        }
        for entry in all_reg_rows:
            row = tbl.add_row().cells
            _ct(row[0], entry.get("product_name", ""))
            _ct(row[1], entry["oem_name"])
            _ct(row[2], entry["regulatory_body"])
            _ct(row[3], entry["region"])
            _ct(row[4], entry["status"])
            bg = _REG_STATUS_COLOURS.get((entry["status"] or "").lower(), "F5F5F5")
            _set_bg(row[4], bg)
            _ct(row[5], entry["confidence"])
        doc.add_paragraph()
        next_sec += 1

    # ── Circular Relationships ───────────────────────────────────────────────
    if metrics.get("circular_pairs"):
        doc.add_heading(f"{next_sec}. Circular Supply Relationships", level=1)
        narr_circ = narrative.get("circular_relationships", "")
        if narr_circ:
            _add_narrative(narr_circ)
        doc.add_paragraph(
            "The following company pairs have a mutual supply relationship "
            "(A supplies B and B supplies A). These warrant further investigation."
        )
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Company A", "Company B", "Note"], "B71C1C")
        for entry in metrics["circular_pairs"]:
            row = tbl.add_row().cells
            label_a = f"⚠ {entry['company_a']}" if entry.get("inferred") else entry["company_a"]
            label_b = f"⚠ {entry['company_b']}" if entry.get("inferred") else entry["company_b"]
            _ct(row[0], label_a)
            _ct(row[1], label_b)
            note = "Model inference" if entry.get("inferred") else "Circular dependency"
            _ct(row[2], note)
            if entry.get("inferred"):
                _set_bg(row[2], "FFFDE7")
            else:
                _set_bg(row[2], "FFEBEE")
        doc.add_paragraph()
        next_sec += 1

    # ── Cross-product overlap ────────────────────────────────────────────────
    if metrics["cross_product_suppliers"]:
        doc.add_heading(f"{next_sec}. Cross-Product Supplier Overlap", level=1)
        _add_narrative(narrative.get("cross_product", ""))
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Company", "Files / Products", "Note"], "E65100")
        for entry in metrics["cross_product_suppliers"]:
            row = tbl.add_row().cells
            _flag_row(row, entry)
            _ct(row[1], ", ".join(entry["products"]))
        doc.add_paragraph()
        next_sec += 1

    # ── Product Breakdown ────────────────────────────────────────────────────
    product_groups = _group_by_source(analysis_rows)
    doc.add_heading(f"{next_sec}. Product Breakdown", level=1)
    doc.add_paragraph(
        "The following sub-sections provide per-product analysis including hub companies, "
        "potential bottlenecks, key components, and supply chain visualisations."
    )
    reg      = regulatory_by_file   or {}
    oem_ctry = oem_countries_by_file or {}

    # Launch one shared Playwright browser for all per-product Leaflet screenshots
    _pw_ctx  = None
    _browser = None
    try:
        from playwright.sync_api import sync_playwright
        _pw_ctx    = sync_playwright().__enter__()
        _browser   = _pw_ctx.chromium.launch(
            headless=True,
            args=["--no-sandbox", "--disable-dev-shm-usage"],
        )
    except Exception:
        pass   # Playwright unavailable — screenshots fall back to None gracefully

    try:
        for sub_i, (display_name, prod_rows) in enumerate(product_groups, 1):
            source_file   = prod_rows[0]["source_file"] if prod_rows else None
            reg_rows_prod = reg.get(source_file, [])
            oem_ctry_prod = oem_ctry.get(source_file, {})
            _add_product_section(doc, _docx_helpers, display_name, prod_rows,
                                 f"{next_sec}.{sub_i}", regulatory_rows=reg_rows_prod,
                                 oem_countries=oem_ctry_prod, browser=_browser)
    finally:
        if _browser:
            try:
                _browser.close()
            except Exception:
                pass
        if _pw_ctx:
            try:
                _pw_ctx.__exit__(None, None, None)
            except Exception:
                pass

    doc.add_paragraph()
    next_sec += 1

    # ── Data Summary ─────────────────────────────────────────────────────────
    doc.add_heading(f"{next_sec}. Data Summary", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _hdr_row(tbl, ["Source File", "Supplier Count"], "37474F")
    for source, count in metrics["tier_distribution"].items():
        row = tbl.add_row().cells
        _ct(row[0], source)
        _ct(row[1], count)
    doc.add_paragraph()
    next_sec += 1

    # ── Requires Further Research ────────────────────────────────────────────
    if research_rows:
        doc.add_heading(f"{next_sec}. Requires Further Research", level=1)
        narr = narrative.get("needs_research", "")
        if narr:
            _add_narrative(narr)
        doc.add_paragraph(
            "The following entries have a confirmed company existence and supply tie, "
            "but the component supplied could not be verified. Follow-up investigation is recommended."
        )
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Company", "Supplies To", "Country", "Source File"], "BF360C")
        for entry in research_rows:
            row = tbl.add_row().cells
            _ct(row[0], entry["company_name"])
            _ct(row[1], entry["supplies_to"])
            _ct(row[2], entry["country"])
            _ct(row[3], entry["source_file"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Background worker ─────────────────────────────────────────────────────────

def _run_insights(stream_id, files_data, per_file_download=True):
    import time as _time
    try:
        (analysis_rows, research_rows, regulatory_by_file,
         oem_countries_by_file, product_info_by_file,
         file_stats_by_file) = _ingest_files(files_data, stream_id)

        if not analysis_rows:
            _push(stream_id, {
                "type": "error",
                "message": "No qualifying rows found. Ensure uploaded files have been through the verification pipeline and have at least one row where Company Exists = Yes and Correct Component Supplied = Yes.",
            })
            return

        total_reg = sum(len(v) for v in regulatory_by_file.values())
        pharma_note = f" 💊 {total_reg} regulatory row(s) found." if total_reg else ""
        _status(stream_id, f"✅ {len(analysis_rows)} analysis rows loaded ({sum(1 for r in analysis_rows if not r['_inferred'])} verified, {sum(1 for r in analysis_rows if r['_inferred'])} inferred). {len(research_rows)} flagged for further research.{pharma_note}")

        # ── Per-file verified supply data reports ─────────────────────────────
        if per_file_download:
            rows_by_file = {}
            for row in analysis_rows:
                rows_by_file.setdefault(row["source_file"], []).append(row)
            research_by_file = {}
            for row in research_rows:
                research_by_file.setdefault(row["source_file"], []).append(row)

            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            for file_idx, filename in enumerate(rows_by_file):
                _status(stream_id, f"📄 Building per-file report for {filename}…")
                pinfo         = product_info_by_file.get(filename, {})
                file_rows     = rows_by_file[filename]
                file_research = research_by_file.get(filename, [])
                file_reg      = regulatory_by_file.get(filename, [])
                per_file_bytes = _build_per_file_docx(
                    pinfo, file_rows, file_research,
                    reg_rows=file_reg or None,
                    oem_countries=oem_countries_by_file.get(filename, {}),
                )

                safe_name = re.sub(r'[^A-Za-z0-9_\-]', '_',
                                   pinfo.get("name") or filename.replace(".xlsx", ""))
                doc_name  = f"{safe_name}_verified_supply_data_report_{ts}.docx"
                file_key  = f"{stream_id}_{file_idx}"

                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(per_file_bytes)
                    tmp_path = tmp.name

                with _insights_file_lock:
                    _insights_file_store[file_key] = {"path": tmp_path, "filename": doc_name}

                _push(stream_id, {"type": "download_file", "key": file_key, "filename": doc_name})
                _time.sleep(0.35)

        # ── Combined insights report ──────────────────────────────────────────
        pandas_metrics = _run_pandas(analysis_rows, stream_id)
        graph_metrics  = _run_networkx(analysis_rows, stream_id)

        products = sorted({r["product"] for r in analysis_rows if r["product"]})
        if not products:
            products = sorted({r["source_file"] for r in analysis_rows})

        metrics = {
            "total_verified_suppliers": sum(1 for r in analysis_rows if not r["_inferred"]),
            "total_inferred_suppliers": sum(1 for r in analysis_rows if r["_inferred"]),
            "products_analysed":        products,
            **pandas_metrics,
            **graph_metrics,
            "needs_research": [
                {
                    "company":     r["company_name"],
                    "supplies_to": r["supplies_to"],
                    "source_file": r["source_file"],
                }
                for r in research_rows
            ],
        }

        narrative = _generate_narrative(metrics, stream_id, regulatory_by_file=regulatory_by_file)

        _status(stream_id, "🗺️ Generating choropleth map…")
        _status(stream_id, "📄 Building combined insights report (per-product charts may take ~10s)…")
        docx_bytes = _build_docx(metrics, narrative, analysis_rows, research_rows,
                                 regulatory_by_file=regulatory_by_file,
                                 oem_countries_by_file=oem_countries_by_file,
                                 file_stats_by_file=file_stats_by_file)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name

        date_str  = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename  = f"supply_chain_insights_{date_str}.docx"
        with _insights_store_lock:
            _insights_store[stream_id] = {"path": tmp_path, "filename": filename}

        _status(stream_id, "✅ Combined report ready — downloading…")

    except Exception as e:
        _push(stream_id, {"type": "error", "message": str(e)})
    finally:
        _push(stream_id, {"type": "stream_end"})


# ── Routes ────────────────────────────────────────────────────────────────────

@insights_bp.route("/dev/insights")
def insights_page():
    html = """<!DOCTYPE html>
<html>
<head>
  <title>Supply Chain Insights</title>
  <link href="https://fonts.googleapis.com/css2?family=Syne:wght@800&display=swap" rel="stylesheet">
  <style>
    .logo{font-family:'Syne',sans-serif;font-weight:800;font-size:1.4rem;letter-spacing:-.02em}
    .logo span{color:#e8ff47}
    body { font-family: monospace; padding: 2rem; background: #111; color: #eee; }
    h2 span { color: #e8ff47; }
    p.desc { color: #aaa; font-size: .88rem; max-width: 700px; line-height: 1.7; margin-bottom: 1.2rem; }
    p.desc strong { color: #eee; }
    button { background: #333; color: #eee; border: 1px solid #555; padding: .5rem 1.5rem;
             cursor: pointer; font-family: monospace; font-size: .9rem; }
    button:hover { background: #444; }
    button:disabled { opacity: .4; cursor: not-allowed; }
    #log-container { height: 420px; overflow-y: auto; background: #1a1a1a; border: 1px solid #333;
                     padding: .75rem; margin-top: 1.5rem; border-radius: 4px; }
    #log div { margin: 2px 0; font-size: .85rem; line-height: 1.6; }
    .done     { color: #4caf50; }
    .error    { color: #f44336; }
    .log-line { color: #888; font-size: .78rem; }
    #timer { display: none; font-size: .85rem; color: #e8ff47; margin-top: .75rem; letter-spacing: .05em; }
    #timer span { font-weight: bold; }
    #file-list { margin: .75rem 0; display: flex; flex-direction: column; gap: .35rem; min-height: 1.5rem; }
    .file-row { display: flex; align-items: center; gap: .75rem; padding: .4rem .65rem; background: #1a1a1a; border: 1px solid #333; border-radius: 3px; }
    .file-name { flex: 1; font-size: .83rem; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
    .file-size { color: #555; font-size: .75rem; flex-shrink: 0; }
    .remove-btn { background: none; border: none; color: #555; cursor: pointer; padding: 0 2px; font-size: .85rem; line-height: 1; flex-shrink: 0; font-family: monospace; }
    .remove-btn:hover:not(:disabled) { color: #f44336; }
    #addBtn { background: #1a2a1a; color: #4caf50; border: 1px dashed #2e5c2e; padding: .35rem 1rem; font-size: .82rem; font-family: monospace; cursor: pointer; }
    #addBtn:hover:not(:disabled) { background: #213321; }
  </style>
</head>
<body>
  <div class="logo" style="margin-bottom:1.5rem">
    <img src="/static/images/nusW-tliap_transparent_bg.png" width="320" height="80" style="vertical-align:middle"> Supplier<span>Map</span>
  </div>
  <h2>Supply Chain <span>Insights</span></h2>
  <p class="desc">
    Upload one or more verified Excel files (.xlsx) exported from the verification pipeline.<br><br>
    <strong>All 3 Yes</strong> — treated as ground truth, used as-is.<br>
    <strong>Company Exists + Correct Component = Yes, Supply Ties unconfirmed</strong> — included but flagged ⚠ as model inference.<br>
    <strong>Company Exists + Supply Ties = Yes, Correct Component unconfirmed</strong> — excluded from analysis, surfaced in a "Requires Further Research" section.<br>
    All other rows are excluded.
  </p>

  <div id="file-list"><div style="color:#555;font-size:.82rem;padding:.3rem 0">No files added.</div></div>
  <input type="file" id="hiddenInput" accept=".xlsx" multiple style="display:none" onchange="addFiles(this.files)">
  <button id="addBtn" onclick="document.getElementById('hiddenInput').click()">+ Add Files</button>

  <div style="margin-top:1rem;display:flex;align-items:center;gap:1.2rem;flex-wrap:wrap">
    <button id="genBtn" onclick="generateInsights()">▶ Generate Report</button>
    <label style="display:flex;align-items:center;gap:.5rem;font-size:.85rem;color:#aaa;cursor:pointer;user-select:none">
      <input type="checkbox" id="perFileChk" checked
             style="width:14px;height:14px;accent-color:#e8ff47;cursor:pointer">
      Download individual verified report for each file
    </label>
  </div>

  <div id="timer">⏱ Elapsed: <span id="timerVal">0.0s</span></div>
  <div id="log-container"><div id="log"></div></div>

  <script>
    const log          = document.getElementById('log');
    const logContainer = document.getElementById('log-container');

    const append = (msg, cls) => {
      const d = document.createElement('div');
      if (cls) d.className = cls;
      d.textContent = msg;
      log.appendChild(d);
      logContainer.scrollTop = logContainer.scrollHeight;
    };

    let fileQueue     = [];
    let timerInterval = null, timerStart = null;

    // ── file queue management ────────────────────────────────────────────────
    function addFiles(newFiles) {
      for (const f of newFiles) {
        if (!fileQueue.some(q => q.name === f.name && q.size === f.size))
          fileQueue.push(f);
      }
      renderFileList();
      document.getElementById('hiddenInput').value = '';
    }

    function removeFile(idx) {
      fileQueue.splice(idx, 1);
      renderFileList();
    }

    function renderFileList() {
      const list = document.getElementById('file-list');
      if (!fileQueue.length) {
        list.innerHTML = '<div style="color:#555;font-size:.82rem;padding:.3rem 0">No files added.</div>';
        return;
      }
      list.innerHTML = fileQueue.map((f, i) => `
        <div class="file-row" id="file-row-${i}">
          <button class="remove-btn" onclick="removeFile(${i})">✕</button>
          <span class="file-name">${f.name}</span>
          <span class="file-size">${(f.size / 1024).toFixed(1)} KB</span>
        </div>`).join('');
    }

    function setRunning(running) {
      document.getElementById('genBtn').disabled  = running;
      document.getElementById('addBtn').disabled  = running;
      document.querySelectorAll('.remove-btn').forEach(b => b.disabled = running);
    }

    // ── timer ────────────────────────────────────────────────────────────────
    function startTimer() {
      timerStart = Date.now();
      document.getElementById('timer').style.display = 'block';
      document.getElementById('timerVal').textContent = '0.0s';
      if (timerInterval) clearInterval(timerInterval);
      timerInterval = setInterval(() => {
        document.getElementById('timerVal').textContent =
          ((Date.now() - timerStart) / 1000).toFixed(1) + 's';
      }, 100);
    }

    function stopTimer() {
      if (timerInterval) { clearInterval(timerInterval); timerInterval = null; }
      if (!timerStart) return;
      document.getElementById('timerVal').textContent =
        ((Date.now() - timerStart) / 1000).toFixed(1) + 's ✓';
    }

    // ── generate ─────────────────────────────────────────────────────────────
    async function generateInsights() {
      if (!fileQueue.length) { append('⚠️ Add at least one .xlsx file first.'); return; }

      log.innerHTML = '';
      setRunning(true);
      append('Uploading ' + fileQueue.length + ' file(s)…');
      startTimer();

      const fd = new FormData();
      for (const f of fileQueue) fd.append('files', f);
      fd.append('per_file_download', document.getElementById('perFileChk').checked ? 'true' : 'false');

      let streamId;
      try {
        const res  = await fetch('/api/insights/generate', { method: 'POST', body: fd });
        const data = await res.json();
        if (data.error) {
          append('Error: ' + data.error, 'error');
          stopTimer(); setRunning(false); return;
        }
        streamId = data.stream_id;
      } catch(e) {
        append('Upload failed: ' + e, 'error');
        stopTimer(); setRunning(false); return;
      }

      append('Files received — running analysis…');
      const es = new EventSource('/api/insights/stream/' + streamId);

      es.onmessage = (e) => {
        const msg = JSON.parse(e.data);
        if (msg.type === 'status') {
          append(msg.message);
        } else if (msg.type === 'log') {
          append(msg.message, msg.is_error ? 'error' : 'log-line');
        } else if (msg.type === 'error') {
          stopTimer(); setRunning(false);
          append('Error: ' + msg.message, 'error');
          es.close();
        } else if (msg.type === 'download_file') {
          const a = document.createElement('a');
          a.href = '/api/insights/download_file/' + msg.key;
          a.download = msg.filename;
          document.body.appendChild(a); a.click(); document.body.removeChild(a);
          append('📄 Downloaded: ' + msg.filename, 'done');
        } else if (msg.type === 'stream_end') {
          stopTimer(); setRunning(false);
          es.close();
          const a = document.createElement('a');
          a.href = '/api/insights/download/' + streamId;
          a.download = '';
          document.body.appendChild(a); a.click(); document.body.removeChild(a);
          append('✓ Combined insights report downloaded.', 'done');
        }
      };

      es.onerror = () => {
        stopTimer(); setRunning(false);
        append('SSE connection lost.', 'error');
        es.close();
      };
    }
  </script>
</body>
</html>"""
    return Response(html, mimetype="text/html")


@insights_bp.route("/api/insights/generate", methods=["POST"])
def insights_generate():
    uploaded = request.files.getlist("files")
    if not uploaded:
        return jsonify({"error": "No files uploaded"}), 400

    files_data = []
    for f in uploaded:
        if not f.filename.lower().endswith(".xlsx"):
            return jsonify({"error": f"{f.filename} is not an .xlsx file"}), 400
        files_data.append({"filename": f.filename, "bytes": f.read()})

    per_file_download = request.form.get("per_file_download", "true").lower() == "true"

    stream_id = uuid.uuid4().hex
    with _insights_queues_lock:
        _insights_queues[stream_id] = []

    threading.Thread(
        target=_run_insights,
        args=(stream_id, files_data, per_file_download),
        daemon=True,
    ).start()

    return jsonify({"stream_id": stream_id})


@insights_bp.route("/api/insights/stream/<stream_id>")
def insights_stream(stream_id):
    import time as _time

    def generate():
        while True:
            with _insights_queues_lock:
                queue  = _insights_queues.get(stream_id, [])
                events, _insights_queues[stream_id] = queue[:], []

            for event in events:
                yield f"data: {json.dumps(event)}\n\n"
                if event.get("type") == "stream_end":
                    with _insights_queues_lock:
                        _insights_queues.pop(stream_id, None)
                    return

            if not events:
                yield ": keepalive\n\n"
            _time.sleep(0.5)

    return Response(generate(), mimetype="text/event-stream", headers={
        "Cache-Control":    "no-cache",
        "X-Accel-Buffering": "no",
    })


@insights_bp.route("/api/insights/download/<stream_id>")
def insights_download(stream_id):
    with _insights_store_lock:
        entry = _insights_store.pop(stream_id, None)

    if not entry:
        return jsonify({"error": "Report not found or already downloaded"}), 404

    path     = entry["path"]
    filename = entry["filename"]

    response = send_file(
        path,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    @response.call_on_close
    def cleanup():
        try:
            os.remove(path)
        except OSError:
            pass

    return response


@insights_bp.route("/api/insights/download_file/<key>")
def insights_download_file(key):
    """Serve a per-file verified supply data report and delete the temp file after."""
    with _insights_file_lock:
        entry = _insights_file_store.pop(key, None)

    if not entry:
        return jsonify({"error": "File report not found or already downloaded"}), 404

    path     = entry["path"]
    filename = entry["filename"]

    response = send_file(
        path,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    @response.call_on_close
    def cleanup():
        try:
            os.remove(path)
        except OSError:
            pass

    return response
