"""
Supplier Chain Mapper — Multi-provider AI backend
Search backends:
  - Gemini provider: uses Gemini's built-in Google Search grounding (no separate search API needed)
  - Anthropic/OpenAI: uses DuckDuckGo (ddgs package)
"""
"""testing branches"""

import json, re, time, threading, os
from dotenv import load_dotenv
load_dotenv()

from flask import Flask, request, jsonify, Response, render_template
from flask_cors import CORS
import requests
from bs4 import BeautifulSoup
import random
import openpyxl
import datetime
import io
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from flask import Response, jsonify, request, send_file
import uuid
import re
import threading
import tempfile

app = Flask(__name__)
CORS(app)

# ── Country coords ────────────────────────────────────────────────────────────
COUNTRY_COORDS = {
    "united states":(37.09,-95.71),"usa":(37.09,-95.71),"us":(37.09,-95.71),
    "china":(35.86,104.19),"prc":(35.86,104.19),
    "japan":(36.20,138.25),"germany":(51.16,10.45),
    "south korea":(35.90,127.76),"korea":(35.90,127.76),
    "taiwan":(23.69,120.96),"united kingdom":(55.37,-3.43),"uk":(55.37,-3.43),
    "france":(46.22,2.21),"india":(20.59,78.96),"canada":(56.13,-106.34),
    "australia":(-25.27,133.77),"brazil":(-14.23,-51.92),"mexico":(23.63,-102.55),
    "netherlands":(52.13,5.29),"sweden":(60.12,18.64),"switzerland":(46.81,8.22),
    "italy":(41.87,12.56),"spain":(40.46,-3.74),"singapore":(1.35,103.81),
    "malaysia":(4.21,108.96),"thailand":(15.87,100.99),"vietnam":(14.05,108.27),
    "indonesia":(-0.78,113.92),"philippines":(12.87,121.77),"russia":(61.52,105.31),
    "saudi arabia":(23.88,45.07),"uae":(23.42,53.84),"united arab emirates":(23.42,53.84),
    "south africa":(-30.55,22.93),"nigeria":(9.08,8.67),"egypt":(26.82,30.80),
    "turkey":(38.96,35.24),"poland":(51.91,19.14),"czech republic":(49.81,15.47),
    "hungary":(47.16,19.50),"finland":(61.92,25.74),"norway":(60.47,8.46),
    "denmark":(56.26,9.50),"belgium":(50.50,4.46),"austria":(47.51,14.55),
    "portugal":(39.39,-8.22),"israel":(31.04,34.85),"new zealand":(-40.90,174.88),
    "argentina":(-38.41,-63.61),"chile":(-35.67,-71.54),"colombia":(4.57,-74.29),
    "bangladesh":(23.68,90.35),"pakistan":(30.37,69.34),"sri lanka":(7.87,80.77),
    "cambodia":(12.56,104.99),"myanmar":(19.15,96.63),
}
# Global Variables for Bulk Exporting (Validation Pipeline)
# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION — edit these
# ─────────────────────────────────────────────────────────────────────────────

"""
"Fentanyl Citrate 800mg Oral Transmucosal",
"Morphine Sulphate 10mg auto injector",
"Pralidoxime and Atropine, 600mg pralidoxime Chloride and at least 2mg Atropine Sulphate, Intramuscular Auto-Injector",
"Obidoxime and Atropine, 220mg Obidoxime Dichloride and 2mg Atropine through Intramuscular Auto-Injector",
"Diazepam 10mg Intramuscular auto-injector",
"Midazolam 10mg intramuscular auto-injector",
"Modafinil tablet 100mg",
"Caffeine tablet 200mg",
"Drone (mini class, weight <=5, size <=450mm) body frame",
"Drone (mini class, weight <=5, size <=450mm) Propeller",
"Drone (mini class, weight <=5, size <=450mm) DC Brushless Motor ",
"Drone (mini class, weight <=5, size <=450mm) Electric Speed Controller",
"Drone (mini class, weight <=5, size <=450mm) Flight Processing Controller",
"Drone (mini class, weight <=5, size <=450mm) Radio Receiver and Transmitter",
"Drone (mini class, weight <=5, size <=450mm) Antenna",
"Drone (mini class, weight <=5, size <=450mm) Camera with Gimbal",
"Drone (mini class, weight <=5, size <=450mm) Ground Controlling Unit",
"Drone (mini class, weight <=5, size <=450mm) GPS",
"Lithium ion battery 18650 cylindrical",
"Lithium ion battery 21700 cylindrical",
"Lithium ion battery 4680 cylindrical",
"Lithium polymer batter pouch cell"
"""

PRODUCTS_TO_TEST = [
    "Fentanyl Citrate 800mg Oral Transmucosal",
    "Morphine Sulphate 10mg auto injector",
    "Pralidoxime and Atropine, 600mg pralidoxime Chloride and at least 2mg Atropine Sulphate, Intramuscular Auto-Injector",
    "Obidoxime and Atropine, 220mg Obidoxime Dichloride and 2mg Atropine through Intramuscular Auto-Injector",
    "Diazepam 10mg Intramuscular auto-injector",
    "Midazolam 10mg intramuscular auto-injector",
    "Modafinil tablet 100mg",
    "Caffeine tablet 200mg",
    "Drone (mini class, weight <=5, size <=450mm) body frame",
    "Drone (mini class, weight <=5, size <=450mm) Propeller",
    "Drone (mini class, weight <=5, size <=450mm) DC Brushless Motor ",
    "Drone (mini class, weight <=5, size <=450mm) Electric Speed Controller",
    "Drone (mini class, weight <=5, size <=450mm) Flight Processing Controller",
    "Drone (mini class, weight <=5, size <=450mm) Radio Receiver and Transmitter",
    "Drone (mini class, weight <=5, size <=450mm) Antenna",
    "Drone (mini class, weight <=5, size <=450mm) Camera with Gimbal",
    "Drone (mini class, weight <=5, size <=450mm) Ground Controlling Unit",
    "Drone (mini class, weight <=5, size <=450mm) GPS",
    "Lithium ion battery 18650 cylindrical",
    "Lithium ion battery 21700 cylindrical",
    "Lithium ion battery 4680 cylindrical",
"Lithium polymer batter pouch cell"
]

DEPTH    = 3       # supply chain depth per product (1–3)
PROVIDER = "gemini" # which provider to use: anthropic | openai | gemini | deepseek


# ─────────────────────────────────────────────────────────────────────────────
# STYLES
# ─────────────────────────────────────────────────────────────────────────────

FONT = "Arial"

def _fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def _font(bold=False, size=9, color="000000"):
    return Font(name=FONT, bold=bold, size=size, color=color)

def _align(wrap=True):
    return Alignment(horizontal="left", vertical="top", wrap_text=wrap)

def _border():
    s = Side(style="thin", color="CCCCCC")
    return Border(left=s, right=s, top=s, bottom=s)

# Header fills per sheet type
HDR_PRODUCT   = _fill("1A237E")   # dark navy  — product overview
HDR_OEM       = _fill("1B5E20")   # dark green — OEM sheet
HDR_TIER      = [
    _fill("1565C0"),   # tier 1 — blue
    _fill("6A1B9A"),   # tier 2 — purple
    _fill("E65100"),   # tier 3 — orange
    _fill("37474F"),   # tier 4 — dark grey
]
ALT_ROW       = _fill("F5F5F5")
WHITE_ROW     = _fill("FFFFFF")

def get_coords(country: str):
    if not country: return None
    key = country.strip().lower()
    return COUNTRY_COORDS.get(key) or COUNTRY_COORDS.get(key.split(",")[0].strip())

def clean_json(text: str) -> str:
    """Strip markdown fences and fix common AI JSON formatting issues."""
    # Remove markdown code fences
    text = re.sub(r"```[a-z]*", "", text).strip("` \n")
    # Remove trailing commas before ] or } (common AI mistake)
    text = re.sub(r",\s*([\]}])", r"\1", text)
    # Remove single-line // comments
    text = re.sub(r"//[^\n]*", "", text)
    # Remove /* */ block comments
    text = re.sub(r"/\*.*?\*/", "", text, flags=re.DOTALL)
    return text.strip()


def extract_balanced(text: str, open_char: str, close_char: str) -> str | None:
    """Extract the largest balanced bracket block from text."""
    best = None
    for start in range(len(text)):
        if text[start] != open_char:
            continue
        depth = 0
        for end in range(start, len(text)):
            if text[end] == open_char:
                depth += 1
            elif text[end] == close_char:
                depth -= 1
                if depth == 0:
                    candidate = text[start:end+1]
                    if best is None or len(candidate) > len(best):
                        best = candidate
                    break
    return best


def safe_parse_json(text: str) -> list | dict | None:
    """Try multiple strategies to parse potentially malformed JSON from AI."""
    cleaned = clean_json(text)

    # Strategy 1: direct parse of cleaned text
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Strategy 2: extract largest balanced [...] array first (preserves all items)
    arr_block = extract_balanced(cleaned, "[", "]")
    if arr_block:
        try:
            return json.loads(arr_block)
        except json.JSONDecodeError:
            # Try fixing trailing commas inside extracted block
            fixed = re.sub(r",\s*([\]}])", r"", arr_block)
            try:
                return json.loads(fixed)
            except json.JSONDecodeError:
                pass

    # Strategy 3: extract largest balanced {...} object
    obj_block = extract_balanced(cleaned, "{", "}")
    if obj_block:
        try:
            return json.loads(obj_block)
        except json.JSONDecodeError:
            fixed = re.sub(r",\s*([\]}])", r"", obj_block)
            try:
                return json.loads(fixed)
            except json.JSONDecodeError:
                pass

    # Strategy 4: close unclosed brackets and retry
    try:
        open_brackets = cleaned.count("[") - cleaned.count("]")
        open_braces   = cleaned.count("{") - cleaned.count("}")
        fixed = cleaned + ("}" * max(0, open_braces)) + ("]" * max(0, open_brackets))
        fixed = re.sub(r",\s*([\]}])", r"", fixed)
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    return None


# ── Search backends ───────────────────────────────────────────────────────────

def ddg_search_and_scrape(query: str, n: int = 4) -> str:
    """
    Try DuckDuckGo search + scraping. If blocked, falls back gracefully.
    Anthropic/OpenAI will use their training knowledge when search fails.
    """
    results = []
    # Try ddgs (new name) then duckduckgo_search (old name)
    for attempt in range(2):
        try:
            time.sleep(2 * (attempt + 1))
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            with DDGS() as ddgs:
                results = list(ddgs.text(query, max_results=n))
            if results:
                break
        except Exception as e:
            print(f"DDG attempt {attempt+1} failed: {e}")

    if not results:
        print(f"Search unavailable for: {query} — AI will use training knowledge")
        return ""   # empty string signals to use knowledge-only prompt

    parts = []
    for r in results[:3]:
        url = r.get("href","")
        snippet = r.get("body","")
        page = ""
        if url:
            try:
                hdrs = {"User-Agent":"Mozilla/5.0 (compatible; SupplierMapper/1.0)"}
                resp = requests.get(url, headers=hdrs, timeout=8)
                resp.raise_for_status()
                soup = BeautifulSoup(resp.text,"html.parser")
                for t in soup(["script","style","nav","footer","header"]): t.decompose()
                page = soup.get_text(separator=" ",strip=True)[:1500]
            except: pass
        parts.append(f"SOURCE: {url}\nSNIPPET: {snippet}\nCONTENT: {page}")
    return "\n---\n".join(parts)

def gemini_search_and_answer(query: str, api_key: str) -> str:
    """
    Use Gemini with Google Search grounding via the new google-genai SDK.
    Gemini searches the web itself — no separate search API key needed.
    """
    try:
        from google import genai
        from google.genai import types
        client = genai.Client(api_key=api_key)
        resp = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=f"Search the web and provide detailed factual information about: {query}\nInclude company names, countries, and supply chain relationships.",
            config=types.GenerateContentConfig(
                tools=[types.Tool(google_search=types.GoogleSearch())]
            )
        )
        # Extract text from response
        text_parts = []
        for part in resp.candidates[0].content.parts:
            if hasattr(part, "text") and part.text:
                text_parts.append(part.text)
        return "\n".join(text_parts) if text_parts else "No results."
    except Exception as e:
        print(f"Gemini search grounding error: {e}, falling back to base knowledge")
        try:
            from google import genai
            client = genai.Client(api_key=api_key)
            resp = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=f"Based on your knowledge, provide information about: {query}\nInclude company names, countries, and supply chain relationships."
            )
            return resp.text
        except Exception as e2:
            return f"Search failed: {e2}"


# ── AI call abstraction ───────────────────────────────────────────────────────

def call_ai(prompt: str, provider: str, max_tokens: int = 1000) -> str:
    if provider == "anthropic":
        import anthropic
        key = os.getenv("ANTHROPIC_API_KEY")
        if not key: raise ValueError("ANTHROPIC_API_KEY not set in .env")
        c = anthropic.Anthropic(api_key=key)
        r = c.messages.create(model="claude-sonnet-4-6", max_tokens=max_tokens,
                               messages=[{"role":"user","content":prompt}])
        return r.content[0].text

    elif provider == "openai":
        from openai import OpenAI
        key = os.getenv("OPENAI_API_KEY")
        if not key: raise ValueError("OPENAI_API_KEY not set in .env")
        c = OpenAI(api_key=key)
        r = c.chat.completions.create(model="gpt-4o-mini", max_tokens=max_tokens,
                                       messages=[{"role":"user","content":prompt}])
        return r.choices[0].message.content

    elif provider == "gemini":
        from google import genai
        key = os.getenv("GEMINI_API_KEY")
        if not key: raise ValueError("GEMINI_API_KEY not set in .env")
        client = genai.Client(api_key=key)
        r = client.models.generate_content(model="gemini-2.0-flash", contents=prompt)
        return r.text

    elif provider == "deepseek":
        from openai import OpenAI
        key = os.getenv("DEEPSEEK_API_KEY")
        if not key: raise ValueError("DEEPSEEK_API_KEY not set in .env")
        client = OpenAI(api_key=key, base_url="https://api.deepseek.com")
        resp = client.chat.completions.create(
            model="deepseek-chat",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return resp.choices[0].message.content

    else:
        raise ValueError(f"Unknown provider: {provider}")


def get_evidence(query: str, provider: str) -> str:
    """Get web evidence using the best method for each provider."""
    if provider == "gemini":
        key = os.getenv("GEMINI_API_KEY","")
        return gemini_search_and_answer(query, key)
    else:
        evidence = ddg_search_and_scrape(query)
        return evidence  # may be empty string if search is unavailable


def build_prompt_with_evidence(base_prompt: str, evidence: str) -> str:
    """
    If evidence is available, inject it into the prompt.
    If not, ask the AI to use its training knowledge instead.
    """
    if evidence and evidence.strip():
        return base_prompt.replace("{EVIDENCE_BLOCK}",
            f"Use this web evidence to inform your answer:\n{evidence}")
    else:
        return base_prompt.replace("{EVIDENCE_BLOCK}",
            "No live web search is available. Use your training knowledge "
            "to provide the most accurate answer you can. Be explicit about "
            "what you know with high vs low confidence.")


# ── Provider status ───────────────────────────────────────────────────────────

def available_providers() -> list[dict]:
    return [
        {"id":"anthropic","name":"Anthropic Claude","model":"claude-sonnet-4-6",
         "env":"ANTHROPIC_API_KEY","configured":bool(os.getenv("ANTHROPIC_API_KEY")),
         "search":"DuckDuckGo"},
        {"id":"openai","name":"OpenAI GPT-4o-mini","model":"gpt-4o-mini",
         "env":"OPENAI_API_KEY","configured":bool(os.getenv("OPENAI_API_KEY")),
         "search":"DuckDuckGo"},
        {"id":"gemini","name":"Google Gemini","model":"gemini-2.0-flash",
         "env":"GEMINI_API_KEY","configured":bool(os.getenv("GEMINI_API_KEY")),
         "search":"Google Search (built-in)"},
        {"id":"deepseek","name":"DeepSeek V3","model":"deepseek-chat",
         "env":"DEEPSEEK_API_KEY","configured":bool(os.getenv("DEEPSEEK_API_KEY")),
         "search":"DuckDuckGo"},
    ]


# ── SSE helper ────────────────────────────────────────────────────────────────

def sse(data: dict) -> str:
    return f"data: {json.dumps(data)}\n\n"


# ── Pipeline ──────────────────────────────────────────────────────────────────

def run_pipeline(product_input: str, depth: int, provider: str, queue: list, oem_context: list, collected_tiers: dict):
    def push(t, **kw): queue.append({"type":t,**kw})

    search_label = "Google Search" if provider == "gemini" else "DuckDuckGo"
    push("status", message=f"🤖 Provider: {provider} | 🔍 Search: {search_label}")
    push("status", message=f"🔍 Identifying product: {product_input}")

    id_evidence = get_evidence(
        f"{product_input} product specification manufacturer supply chain", provider)

    evidence_note = (
        f"Web research findings:\n{id_evidence}" if id_evidence and id_evidence.strip()
        else "No live web search available — use your training knowledge."
    )
    search_status = "✅ Web evidence gathered" if id_evidence and id_evidence.strip() else "⚠️ Search unavailable — using AI training knowledge"
    push("status", message=search_status)
    

    id_prompt = f"""You are a supply chain research assistant.
Product/Part input: "{product_input}"

{evidence_note}

Return ONLY valid JSON (no markdown) with this exact schema:
{{
  "product_name": "...",
  "product_category": "...",
  "industry": "...",
  "description": "...",
  "key_components": ["...", "..."],
  "oem_manufacturer": "... or unknown",
  "oem_country": "country name or unknown"
}}"""

    try:
        raw = call_ai(id_prompt, provider, max_tokens=800)
        parsed_product = safe_parse_json(raw)
        product_info = parsed_product if isinstance(parsed_product, dict) else None
        if not product_info: raise ValueError("No valid JSON returned")
    except Exception as e:
        push("status", message=f"⚠️ Product ID parse error: {e}")
        product_info = {
            "product_name": product_input, "product_category":"Unknown",
            "industry":"Unknown", "description":"Could not auto-identify.",
            "key_components":[], "oem_manufacturer":"Unknown","oem_country":"Unknown"
        }

    coords = get_coords(product_info.get("oem_country",""))
    if coords: product_info["lat"], product_info["lng"] = coords

    push("product_identified", product=product_info)
    push("status", message=f"✅ Identified: {product_info.get('product_name')}")

    # ── OEM Discovery: find ALL manufacturers of this product ─────────────────
    push("status", message="🏢 Searching for all OEM manufacturers of this product…")

    oem_evidence = get_evidence(
        f"{product_info.get('product_name')} manufacturers OEM brands companies who make produce", provider)
    oem_note = (
        "Web research findings:\n" + oem_evidence if oem_evidence and oem_evidence.strip()
        else "No live web search — use your training knowledge."
    )

    existing_note = ""
    if oem_context:
        names = ", ".join(oem_context)
        existing_note = f"""
    EXISTING COMPANIES ALREADY IN LIST: {names}
    - If you know any of these companies by a different name, use EXACTLY the name shown above
    - Do not add duplicates — only add companies genuinely not in this list
    """



    oem_prompt = f"""You are a supply chain research assistant.
Product: {product_info.get('product_name')} ({product_info.get('industry')})

{oem_note}

Find ALL known OEM manufacturers / brands that produce or sell this product or equivalent products.
Include the primary OEM already identified ({product_info.get('oem_manufacturer')}) plus any others.

CRITICAL COMPANY NAME RULES:
- Use the shortest globally recognised name only (e.g. "Samsung" not "Samsung Electronics Co. Ltd.")
- No legal suffixes: drop Inc, Ltd, LLC, GmbH, Co., Corp, Group, Holdings
- Use English names only (e.g. "Panasonic" not "Panasonic Corporation")
- Be consistent: if a company is known by an acronym (TSMC, BASF, ABB), use the acronym
- If the name of the company is followed by (), check the contents within the bracket, and if it is another name for the said company, drop the brackets and its contents

{existing_note}


CRITICAL: Return ONLY a raw JSON array starting with [ and ending with ].
Do NOT return a single object. Do NOT wrap in markdown. Do NOT add explanation text.
Each OEM must be a SEPARATE element in the array.

Example of correct format:
[
  {{"company_name": "Company A", "country": "Japan", "role": "OEM manufacturer", "market_share": "high", "confidence": "high", "notes": "..."}},
  {{"company_name": "Company B", "country": "China", "role": "Contract manufacturer", "market_share": "medium", "confidence": "high", "notes": "..."}}
]

Schema for each element:
{{
  "company_name": "exact company name",
  "country": "country name",
  "role": "OEM manufacturer | Brand owner | Contract manufacturer | Licensor",
  "market_share": "high|medium|low|unknown",
  "confidence": "high|medium|low",
  "notes": "brief note"
}}

Confidence: high=publicly documented, medium=limited docs, low=inferred.
List 2-8 significant OEMs as SEPARATE array elements."""

    try:
        raw_oem = call_ai(oem_prompt, provider, max_tokens=1000)
        print(f"\n[OEM RAW OUTPUT]:\n{raw_oem[:800]}\n")  # debug
        parsed = safe_parse_json(raw_oem)
        print(f"[OEM PARSED TYPE]: {type(parsed)}, value: {str(parsed)[:200]}")  # debug
        # Handle case where AI returns a dict with a nested list
        if isinstance(parsed, dict):
            for v in parsed.values():
                if isinstance(v, list):
                    parsed = v
                    break
        oem_list = parsed if isinstance(parsed, list) else []
        # If still empty or single item, try splitting by newline JSON objects
        if len(oem_list) <= 1:
            # Try extracting multiple {...} objects from raw text
            objects = re.findall(r'\{[^{}]+\}', raw_oem, re.DOTALL)
            if len(objects) > len(oem_list):
                recovered = []
                for obj in objects:
                    try:
                        recovered.append(json.loads(obj))
                    except:
                        try:
                            recovered.append(json.loads(re.sub(r',\s*([\]}])', r'', obj)))
                        except:
                            pass
                if len(recovered) > len(oem_list):
                    oem_list = recovered
                    push("status", message=f"  🔧 Recovered {len(oem_list)} OEMs from raw text")
        push("status", message=f"  📋 Parsed {len(oem_list)} OEM(s)")
    except Exception as e:
        push("status", message=f"⚠️ OEM discovery parse error: {e}")
        oem_list = []

    # Geocode OEMs
    for oem in oem_list:
        name = oem.get("company_name", "")
        if name and name not in oem_context:
            oem_context.append(name)
        c = get_coords(oem.get("country",""))
        if c: oem["lat"], oem["lng"] = c

    # Ensure primary OEM is always in the list
    primary_oem = product_info.get("oem_manufacturer","")
    if primary_oem and primary_oem.lower() not in ("unknown",""):
        existing_names = [o.get("company_name","").lower() for o in oem_list]
        if not any(primary_oem.lower() in n or n in primary_oem.lower() for n in existing_names):
            primary_coords = get_coords(product_info.get("oem_country",""))
            primary_entry = {
                "company_name": primary_oem,
                "country": product_info.get("oem_country",""),
                "role": "OEM manufacturer",
                "market_share": "high",
                "confidence": "high",
                "notes": "Primary identified manufacturer"
            }
            if primary_coords:
                primary_entry["lat"], primary_entry["lng"] = primary_coords
            oem_list.insert(0, primary_entry)

    push("oem_discovered", oems=oem_list)
    push("status", message=f"✅ Found {len(oem_list)} OEM manufacturer(s)")

    supply_chain = {"product": product_info, "oems": oem_list, "tiers":{}}

    # Build per-OEM supplier map: each OEM searched individually
    # previous_parents: list of {name, oem_root, confidence} dicts
    oem_names = [{"name": o.get("company_name",""), "confidence": o.get("confidence", "")} 
                 for o in oem_list if o.get("company_name","").lower() not in ("unknown","")]
    if not oem_names:
        oem_names = [{"name": product_info.get("oem_manufacturer", product_input), "confidence": "high"}]

    # For tier 1: each parent IS an OEM; for tier 2+: parents are tier-1 suppliers
    # We track parent → oem_root so results stay grouped
    # Structure: list of {"name": str, "oem_root": str}
    previous_parents = [{"name": n.get("name", ""), "oem_root": n.get("name", ""), "confidence": n.get("confidence", "")} for n in oem_names]
    def tier_note(tier_num):
        key = f"tier_{tier_num}"
        names = collected_tiers.get(key, [])
        return f"The following companies are already in the tier {tier_num} list: {', '.join(names)}. For each of the company names, check if it is an alternate name for any of the names in the given list. If you find the same company, use EXACTLY the same name as in the list." if names else ""
    
    for tier_num in range(1, depth + 1):
        push("status", message=f"🏭 Researching Tier-{tier_num} suppliers…")
        tier_suppliers = []   # all suppliers this tier across all parents
        next_parents   = []   # feeds into next tier loop
        # If need to sort or ensure got all 3 tiers of confidence, sort here before looping
        for parent_info in previous_parents[:6]:  # max 6 parents per tier
            parent    = parent_info["name"]
            oem_root  = parent_info["oem_root"]
            if not parent or parent.lower() == "unknown":
                continue

            query = (
                f"{parent} direct Tier 1 suppliers manufacturers components parts"
                if tier_num == 1 else
                f"{parent} suppliers raw materials components manufacturers supply chain"
            )

            evidence = get_evidence(query, provider)
            evidence_note = (
                f"Web research findings:\n{evidence}" if evidence and evidence.strip()
                else "No live web search — use your training knowledge."
            )
            status_icon = "🌐" if evidence and evidence.strip() else "🧠"
            push("status", message=f"  {status_icon} [{oem_root}] → {parent}: searching suppliers…")

            tier_prompt = f"""You are a supply chain research assistant.
Product: {product_info.get('product_name')} ({product_info.get('industry')})
OEM root: {oem_root}
Finding Tier-{tier_num} suppliers — companies that DIRECTLY supply components to: {parent}

{evidence_note}

CRITICAL COMPANY NAME RULES:
- Use the shortest globally recognised name only (e.g. "Samsung" not "Samsung Electronics Co. Ltd.")
- No legal suffixes: drop Inc, Ltd, LLC, GmbH, Co., Corp, Group, Holdings
- Use English names only (e.g. "Panasonic" not "Panasonic Corporation")
- Be consistent: if a company is known by an acronym (TSMC, BASF, ABB), use the acronym
- If the name of the company is followed by (), check the contents within the bracket, and if it is another name for the said company, drop the brackets and its contents
{tier_note(tier_num)}

Return ONLY a valid JSON array (no markdown). Each element:
{{
  "company_name": "exact company name",
  "country": "country name",
  "supplies_to": "{parent}",
  "oem_root": "{oem_root}",
  "components_supplied": ["component1", "component2"],
  "confidence": "high|medium|low",
  "source_hint": "brief source note"
}}

List 3-5 distinct real direct suppliers to {parent} specifically. Do NOT mix in suppliers of other companies."""

            try:
                raw = call_ai(tier_prompt, provider, max_tokens=1200)
                found = safe_parse_json(raw)
                if isinstance(found, list):
                    for s in found:
                        s["oem_root"] = oem_root   # ensure tag is set
                        s["supplies_to"] = parent  # ensure correct parent
                    tier_suppliers.extend(found)
                    # Register these as parents for next tier
                    for s in found:
                        sname = s.get("company_name","").strip()
                        if sname and sname.lower() != "unknown":
                            # added confidence level so can sort if in production, looking for more variety
                            next_parents.append({"name": sname, "oem_root": oem_root, "confidence": s.get("confidence")})
                    previous_parents = next_parents # assign next_parents back to loop
            except Exception as e:
                push("status", message=f"  ⚠️ Parse error tier {tier_num} [{parent}]: {e}")

        # Deduplicate by (company_name + oem_root) — same company can appear under different OEMs
        seen, unique = set(), []
        for s in tier_suppliers:
            #key = f"{s.get('company_name','').strip().lower()}|{s.get('oem_root','')}"
            key = s.get('company_name', '').strip().lower()
            if key and key not in seen:
                seen.add(key)
                c = get_coords(s.get("country",""))
                if c: s["lat"], s["lng"] = c
                unique.append(s)
        for s in unique:
            name = s.get('company_name', '')
            collected_tiers.setdefault(f"tier_{tier_num}", [])
            if name and name not in collected_tiers[f"tier_{tier_num}"]:
                collected_tiers[f"tier_{tier_num}"].append(name)

        supply_chain["tiers"][f"tier_{tier_num}"] = unique
        push("tier_complete", tier=tier_num, suppliers=unique)
        push("status", message=f"✅ Tier-{tier_num}: {len(unique)} suppliers found")

        previous_tier_suppliers = [s.get("company_name","") for s in unique]
        if not previous_tier_suppliers:
            push("status", message=f"⚠️ No suppliers at Tier-{tier_num}, stopping.")
            break

    push("status", message="📊 Generating executive summary…")
    summary_prompt = f"""You are a supply chain analyst.
Supply chain data:
{json.dumps(supply_chain, indent=2)[:6000]}

Write a concise executive summary (150-200 words) covering:
- What the product is and its industry
- Key Tier-1 suppliers and their roles
- Geographic concentration or risks
- Overall supply chain complexity
Return plain text only, no markdown."""

    try:
        supply_chain["summary"] = call_ai(summary_prompt, provider, max_tokens=400).strip()
    except Exception as e:
        supply_chain["summary"] = f"Summary unavailable: {e}"
    
    # added provider key to supply_chain JSON object
    supply_chain["provider"] = provider

    push("complete", supply_chain=supply_chain)

# ── Validation Export Helpers ─────────────────────────────────────────────────

# ── HELPERS ────────────────────────────────────────────────────────────────────

def set_col_widths(ws, widths: list[int]):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

def hdr_cell(ws, row, col, value, fill):
    c = ws.cell(row=row, column=col, value=value)
    c.fill = fill
    c.font = _font(bold=True, size=9, color="FFFFFF")
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    c.border = _border()
    return c

def data_cell(ws, row, col, value, fill=None, bold=False):
    c = ws.cell(row=row, column=col, value=value)
    c.font = _font(bold=bold)
    c.alignment = _align()
    c.border = _border()
    if fill:
        c.fill = fill
    return c

def _row_fill(idx):
    return ALT_ROW if idx%2 == 0 else WHITE_ROW

# ── SHEET WRITERS ────────────────────────────────────────────────────────────────────

def write_index_sheet(ws, results: list[dict], run_meta: dict):
    """
    First sheet: run metadata + one summary row per product
    """
    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 45

    meta_rows = [
        ("Run Date",     run_meta["run_date"]),
        ("Provider",     run_meta["provider"]),
        ("Depth",        str(run_meta["depth"])),
        ("Products Run", str(run_meta["total_products"])),
        ("Total Suppliers Collected", str(run_meta["total_suppliers"])),
    ]
    title = ws.cell(row=1, column=1, value="Bulk Export - Run Summary")
    title.font = Font(name=FONT, bold=True, size=14, color="1A237E")
    ws.merge_cells("A1:B1")

    for r_idx, (k, v) in enumerate(meta_rows, 3):
        data_cell(ws, r_idx, 1, k, bold=True, fill=_fill("E8EAF6"))
        data_cell(ws, r_idx, 2, v)

    # Per-product summary
    tbl_start = len(meta_rows) + 5
    headers = ["Product Input", "Identified Name", "Industry",
               "OEMs Found", "T1 Suppliers", "T2 Suppliers", "T3 Suppliers",
               "Sheet Name"]
    col_widths = [28, 28, 18, 10, 12, 12, 12, 22]
    set_col_widths(ws, col_widths)

    for ci, h in enumerate(headers, 1):
        hdr_cell(ws, tbl_start, ci, h, HDR_PRODUCT)
    ws.row_dimensions[tbl_start].height = 24

    for ri, r in enumerate(results, tbl_start + 1):
        sc = r.get("supply_chain", {})
        fill = _row_fill(ri)
        data_cell(ws, ri, 1, r["product_input"], fill=fill, bold=True)
        data_cell(ws, ri, 2, sc.get("product", {}).get("product_name","—"), fill=fill)
        data_cell(ws, ri, 3, sc.get("product", {}).get("industry","—"), fill=fill)
        data_cell(ws, ri, 4, len(sc.get("oems", [])), fill=fill)
        data_cell(ws, ri, 5, len(sc.get("tiers", {}).get("tier_1",[])), fill=fill)
        data_cell(ws, ri, 6, len(sc.get("tiers", {}).get("tier_2",[])), fill=fill)
        data_cell(ws, ri, 7, len(sc.get("tiers", {}).get("tier_3",[])), fill=fill)
        data_cell(ws, ri, 8, r.get("sheet_prefix",""), fill=fill)

    ws.freeze_panes = f"A{tbl_start + 1}"
    ws.auto_filter.ref = (
        f"A{tbl_start}:{get_column_letter(len(headers))}{tbl_start + len(results)}"
    )

def write_oem_sheet(ws, results: list[dict]):
    """
    One consolidated OEM sheet across all products.
    Columns: Product | Company | Country | Role | Market Share | Confidence | Notes
    """
    headers = ["Product", "Company Name", "Country", "Role", "Market Share", "Confidence", "AI Provider", "Notes"]
    col_widths = [28, 26, 14, 22, 12, 11, 35]
    set_col_widths(ws, col_widths)

    for ci, h in enumerate(headers, 1):
        hdr_cell(ws, 1, ci, h, HDR_OEM)
    ws.row_dimensions[1].height = 24

    row = 2
    for r in results:
        for prov_id, prov_sc in r.get("provider_results", {r.get("provider","unknown"): r.get("supply_chain", {})}).items():
            product_name = prov_sc.get("product", {}).get("product_name", r["product_input"])
            for oem in prov_sc.get("oems", []):
                fill = _row_fill(row)
                data_cell(ws, row, 1, product_name, bold=True, fill=fill)
                data_cell(ws, row, 2, oem.get("company_name", ""), fill=fill)
                data_cell(ws, row, 3, oem.get("country", ""), fill=fill)
                data_cell(ws, row, 4, oem.get("role", ""), fill=fill)
                data_cell(ws, row, 5, oem.get("market_share", ""), fill=fill)
                data_cell(ws, row, 6, oem.get("confidence", ""), fill=fill)
                data_cell(ws, row, 7, oem.get("notes", ""), fill=fill)
                data_cell(ws, row, 8, prov_id, fill=fill)
                row += 1
    
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{row - 1}"

def write_tier_sheet(ws, results: list[dict], tier_key: str, tier_num: int):
    """
    One sheet per tier, all products combined.
    Columns: Product | Company | Country | Supplies To | OEM Root | Components | Confidence | Source
    """
    headers = ["Product", "Company Name", "Country", "Supplies To", "OEM Root", "Components Supplied", "Confidence", "Source Hint", "AI Provider"]
    col_widths = [28, 26, 14, 22, 20, 32, 11, 35]
    set_col_widths(ws, col_widths)

    hdr_fill = HDR_TIER[tier_num-1] if tier_num - 1 < len(HDR_TIER) else _fill("37474F")
    for ci, h in enumerate(headers, 1):
        hdr_cell(ws, 1, ci, h, hdr_fill)
    ws.row_dimensions[1].height = 24

    row = 2
    for r in results:
        for prov_id, prov_sc in r.get("provider_results", {r.get("provider","unknown"): r.get("supply_chain", {})}).items():
            product_name = prov_sc.get("product", {}).get("product_name", r["product_input"])
            suppliers = prov_sc.get("tiers", {}).get(tier_key, [])
            for s in suppliers:
                fill = _row_fill(row)
                components = ", ".join(s.get("components_supplied") or [])
                data_cell(ws, row, 1, product_name, bold=True, fill=fill)
                data_cell(ws, row, 2, s.get("company_name",""), fill=fill)
                data_cell(ws, row, 3, s.get("country", ""), fill=fill)
                data_cell(ws, row, 4, s.get("supplies_to", ""), fill=fill)
                data_cell(ws, row, 5, s.get("oem_root", ""), fill=fill)
                data_cell(ws, row, 6, components, fill=fill)
                data_cell(ws, row, 7, s.get("confidence", ""), fill=fill)
                data_cell(ws, row, 8, s.get("source_hint", ""), fill=fill)
                data_cell(ws, row, 9, prov_id, fill=fill)
                row += 1
    
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{row-1}"


def write_product_sheet(ws, result: dict):
    provider_results = result.get("provider_results", {})
    
    # Fallback: wrap supply_chain as a single-provider result if no provider_results
    if not provider_results:
        sc = result.get("supply_chain", {})
        provider_results = {sc.get("provider", "unknown"): sc}

    # Use the first result's product info for the sheet title/overview
    first_sc = next(iter(provider_results.values()), {})
    product = first_sc.get("product", {})

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 70

    title = ws.cell(row=1, column=1, value=product.get("product_name", result["product_input"]))
    title.font = Font(name=FONT, bold=True, size=13, color="1A237E")
    ws.merge_cells("A1:B1")

    overview_fields = [
        ("Category",        product.get("product_category", "")),
        ("Industry",        product.get("industry", "")),
        ("OEM Manufacturer",product.get("oem_manufacturer", "")),
        ("OEM Country",     product.get("oem_country", "")),
        ("Key Components",  ", ".join(product.get("key_components") or [])),
        ("Description",     product.get("description", "")),
    ]
    for r_idx, (label, value) in enumerate(overview_fields, 3):
        data_cell(ws, r_idx, 1, label, bold=True, fill=_fill("E8EAF6"))
        data_cell(ws, r_idx, 2, value)
        ws.row_dimensions[r_idx].height = 90

    current_row = len(overview_fields) + 5

    tier_headers = ["Company Name", "Country", "Supplies To", "OEM Root", "Components", "Confidence", "Source Hint"]
    tier_widths  = [28, 14, 22, 20, 32, 11, 35]

    # ── One section per provider ──────────────────────────────────────────────
    for prov_id, sc in provider_results.items():

        # Provider heading
        prov_heading = ws.cell(row=current_row, column=1, value=f"Provider: {prov_id.upper()}")
        prov_heading.font = Font(name=FONT, bold=True, size=12, color="FFFFFF")
        prov_heading.fill = _fill("37474F")
        ws.merge_cells(f"A{current_row}:G{current_row}")
        ws.row_dimensions[current_row].height = 22
        current_row += 1

        # Executive summary for this provider
        summary_cell = ws.cell(row=current_row, column=1, value=sc.get("summary", "No summary generated."))
        summary_cell.font = _font()
        summary_cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        ws.merge_cells(f"A{current_row}:G{current_row}")
        ws.row_dimensions[current_row].height = 50
        current_row += 2

        for ti, (tier_key, suppliers) in enumerate(sc.get("tiers", {}).items()):
            tier_num  = ti + 1
            hdr_fill  = HDR_TIER[ti] if ti < len(HDR_TIER) else _fill("37474F")
            tier_label = tier_key.replace("_", " ").upper()

            # Set col widths once
            if ti == 0 and prov_id == next(iter(provider_results)):
                for ci, w in enumerate(tier_widths, 1):
                    ws.column_dimensions[get_column_letter(ci)].width = w

            # Tier heading
            heading = ws.cell(row=current_row, column=1,
                               value=f"{tier_label} ({len(suppliers)} suppliers)")
            heading.font = Font(name=FONT, bold=True, size=11, color="FFFFFF")
            heading.fill = hdr_fill
            ws.merge_cells(f"A{current_row}:G{current_row}")
            ws.row_dimensions[current_row].height = 22
            current_row += 1

            # Column headers
            for ci, h in enumerate(tier_headers, 1):
                hdr_cell(ws, current_row, ci, h, hdr_fill)
            ws.row_dimensions[current_row].height = 20
            current_row += 1

            if not suppliers:
                ws.cell(row=current_row, column=1, value="No suppliers identified")
                current_row += 2
                continue

            for s in suppliers:
                fill = _row_fill(current_row)
                data_cell(ws, current_row, 1, s.get("company_name", ""), bold=True, fill=fill)
                data_cell(ws, current_row, 2, s.get("country", ""),       fill=fill)
                data_cell(ws, current_row, 3, s.get("supplies_to", ""),   fill=fill)
                data_cell(ws, current_row, 4, s.get("oem_root", ""),      fill=fill)
                data_cell(ws, current_row, 5, ", ".join(s.get("components_supplied") or []), fill=fill)
                data_cell(ws, current_row, 6, s.get("confidence", ""),    fill=fill)
                data_cell(ws, current_row, 7, s.get("source_hint", ""),   fill=fill)
                ws.row_dimensions[current_row].height = 60
                current_row += 1

            current_row += 1  # spacing between tiers

        current_row += 2  # spacing between providers

# ─────────────────────────────────────────────────────────────────────────────
# WORKBOOK BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def build_bulk_workbook(results: list[dict], provider: str, depth: int) -> bytes:

    wb = openpyxl.Workbook()
    total_suppliers = sum(sum(len(r["supply_chain"].get("tiers", {}).get(f"tier_{t}", []))
                              for t in range(1, 4))
                              for r in results)
    
    run_meta = {
        "run_date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M UTC"),
        "provider": provider,
        "depth": depth,
        "total_products": len(results),
        "total_suppliers": total_suppliers
    }

    # Index sheet
    ws_index = wb.active
    ws_index.title = "INDEX"
    write_index_sheet(ws_index, results, run_meta)

    # All OEMs sheet
    ws_oem = wb.create_sheet("ALL_OEMs")
    write_oem_sheet(ws_oem, results)

    # Tier 1 - Tier 3 consolidated sheets
    for tier_num in range(1, depth + 1):
        tier_key = f"tier_{tier_num}"
        ws_tier = wb.create_sheet(f"ALL_TIER_{tier_num}")
        write_tier_sheet(ws_tier, results, tier_key, tier_num)

    # Per product detail sheets
    for r in results:
        ws_prod = wb.create_sheet(r["sheet_prefix"])
        write_product_sheet(ws_prod, r)
    
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()





# ── Routes ────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/test')
def test():
    print("testing in app.py")
    return {"test": "ok"}

@app.route('/dev/bulk_export')
def bulk_export():
    collected_oems = []
    collected_tiers = {}
    def safe_sheet_name(product_input: str, index: int) -> str:
        """
        Excel sheet names: max 31 chars, no special characters.
        Prefix with index so identical names don't collide.
        """
        import re
        clean = re.sub(r'[\\/*?\[\]:]', '', product_input)
        clean = clean.strip()[:26]
        return f"{index}_{clean}"
    
    providers = available_providers()
    products = PRODUCTS_TO_TEST
    results = []
    for i, product_input in enumerate(products, 1):
        print(f"\n[bulk_export] ── Product {i}/{len(products)}: {product_input}")

        # Running pipeline synchronously in list queue
        queue = []
        try:
            for provider in providers:
                print(f"\n[bulk_export] - Running pipeline for {provider.get("name")}")
                run_pipeline(product_input, DEPTH, provider.get("id"), queue, collected_oems, collected_tiers)
        except Exception as e:
            print(f"[bulk_export] Pipeline error for '{product_input}': {e}")
            results.append({
                "product_input": product_input,
                "supply_chain":  {"product": {"product_name": product_input},
                                    "oems": [], "tiers": {}, "summary": f"Error: {e}"},
                "sheet_prefix":  safe_sheet_name(product_input, i),
            })
            continue

        # Extracting supply_chains from "complete" event
        supply_chain = None
        for event in queue:
            if event.get("type") == "complete":
                incoming = event.get("supply_chain", {})
                provider_name = incoming.get("provider", "unknown")
                if supply_chain is None:
                    supply_chain = incoming
                else:
                    # merge oems
                    existing_oem_names = {o.get("company_name","").lower() for o in supply_chain.get("oems",[])}
                    for oem in incoming.get("oems", []):
                        if oem.get("company_name","").lower() not in existing_oem_names:
                            supply_chain["oems"].append(oem)

                    # merge tiers
                    for tier_key, suppliers in incoming.get("tiers", {}).items():
                        if tier_key not in supply_chain["tiers"]:
                            supply_chain["tiers"][tier_key] = suppliers
                        else:
                            existing_names = {s.get("company_name","").lower() for s in supply_chain["tiers"][tier_key]}
                            for s in suppliers:
                                if s.get("company_name","").lower() not in existing_names:
                                    supply_chain["tiers"][tier_key].append(s)

        if not supply_chain:
            supply_chain = {"product": {"product_name": product_input}, "oems": [], "tiers": {}}
        
        results.append({
            "product_input": product_input,
            "supply_chain": supply_chain,
            "sheet_prefix": safe_sheet_name(product_input, i)
        })

        tier_counts = {k: len(v) for k, v in supply_chain.get("tiers", {}).items()}
        print(f"[bulk_export] Done: oems={len(supply_chain.get('oems', []))}")
        print(f"tiers={tier_counts}")

    print(f"\n[bulk_export] Building workbook for {len(results)} products...")
    wb_bytes = build_bulk_workbook(results, "All Providers", DEPTH)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"bulk_export_{timestamp}.xlsx"
    print(f"[bulk_export] Sending {len(wb_bytes)} bytes as {filename}")

    return Response(
        wb_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

# multi-single file exports test
_export_files = {}
_export_files_lock = threading.Lock()

def safe_filename(product_input: str, index: int) -> str:
    clean = re.sub(r'[\\/*?:"<>|]', '', product_input).strip()[:50]
    return f"{index:02d}_{clean}.xlsx"

def safe_sheet_name(product_input: str, index: int) -> str:
    clean = re.sub(r'[\\/*?\[\]:]', '', product_input).strip()[:26]
    return f"{index}_{clean}"


@app.route('/dev/multi_single_export')
def multi_single_export():
    accept = request.headers.get("Accept", "")
    if "text/html" in accept:
        return Response("""<!DOCTYPE html>
<html>
<head>
  <title>Bulk Export</title>
  <style>
    body { font-family: monospace; padding: 2rem; background: #111; color: #eee; }
    .done { color: #4caf50; }
    .error { color: #f44336; }
  </style>
</head>
<body>
  <h2>Bulk Export</h2>
  <div id="log"></div>
  <script>
    const log = document.getElementById('log');
    const append = (msg, cls) => {
      const d = document.createElement('div');
      d.className = cls || '';
      d.textContent = msg;
      log.appendChild(d);
    };

    const evtSource = new EventSource('/dev/multi_single_export/stream');

    evtSource.onmessage = (e) => {
      const msg = JSON.parse(e.data);
      if (msg.type === 'start') {
        append(`Starting export for ${msg.total} products...`);
      } else if (msg.type === 'file_ready') {
        append(`[${msg.index}/${msg.total}] ✓ ${msg.filename}`, 'done');
        const a = document.createElement('a');
        a.href = `/dev/multi_single_export/download/${msg.file_id}`;
        a.download = msg.filename;
        document.body.appendChild(a);
        a.click();
        document.body.removeChild(a);
      } else if (msg.type === 'error') {
        append(`[error] ${msg.product}: ${msg.error}`, 'error');
      } else if (msg.type === 'done') {
        append(`All ${msg.total} exports complete.`, 'done');
        evtSource.close();
      }
    };

    evtSource.onerror = () => {
      append('SSE connection lost.', 'error');
      evtSource.close();
    };
  </script>
</body>
</html>""", mimetype="text/html")

    return Response(status=400)


@app.route('/dev/multi_single_export/stream')
def multi_export_stream():
    def generate():
        collected_oems = []
        collected_tiers = {}
        providers = available_providers()
        products = PRODUCTS_TO_TEST
        total = len(products)

        yield f"data: {json.dumps({'type': 'start', 'total': total})}\n\n"

        for i, product_input in enumerate(products, 1):
            print(f"\n[bulk_export] ── Product {i}/{total}: {product_input}")

            queue = []
            try:
                for provider in providers:
                    if not provider.get("configured"):
                        print(f"[bulk_export] Skipping {provider.get('name')} — no API key")
                        continue
                    print(f"[bulk_export] - Running pipeline for {provider.get('name')}")
                    run_pipeline(product_input, DEPTH, provider.get("id"), queue, collected_oems, collected_tiers)
            except Exception as e:
                print(f"[bulk_export] Pipeline error for '{product_input}': {e}")
                yield f"data: {json.dumps({'type': 'error', 'product': product_input, 'error': str(e)})}\n\n"
                continue

            # Merge supply chains from all providers
            provider_results = {}
            for event in queue:
                if event.get("type") == "complete":
                    sc = event.get("supply_chain", {})
                    prov = sc.get("provider", "unknown")
                    if sc.get("oems") or sc.get("tiers"):  # only store if it has data
                        provider_results[prov] = sc

            # derive a merged supply_chain for the index sheet counts
            supply_chain = {"product": {}, "oems": [], "tiers": {}}
            for sc in provider_results.values():
                if not supply_chain["product"]:
                    supply_chain["product"] = sc.get("product", {})
                for oem in sc.get("oems", []):
                    existing = [o.get("company_name","").lower() for o in supply_chain["oems"]]
                    if oem.get("company_name","").lower() not in existing:
                        supply_chain["oems"].append(oem)
                for tier_key, suppliers in sc.get("tiers", {}).items():
                    supply_chain["tiers"].setdefault(tier_key, [])
                    existing = [s.get("company_name","").lower() for s in supply_chain["tiers"][tier_key]]
                    for s in suppliers:
                        if s.get("company_name","").lower() not in existing:
                            supply_chain["tiers"][tier_key].append(s)

            if not supply_chain["product"]:
                supply_chain["product"] = {"product_name": product_input}

            # Build workbook for this product
            result = [{
                "product_input": product_input,
                "supply_chain": supply_chain,
                "provider_results": provider_results,   
                "sheet_prefix": safe_sheet_name(product_input, i)
            }]
            wb_bytes = build_bulk_workbook(result, "All Providers", DEPTH)

            # Save to temp file
            file_id = str(uuid.uuid4())
            filename = safe_filename(product_input, i)
            tmp_path = os.path.join(tempfile.gettempdir(), f"{file_id}.xlsx")

            with open(tmp_path, "wb") as f:
                f.write(wb_bytes)

            with _export_files_lock:
                _export_files[file_id] = {"path": tmp_path, "filename": filename}

            tier_counts = {k: len(v) for k, v in supply_chain.get("tiers", {}).items()}
            print(f"[bulk_export] Ready: {filename} | oems={len(supply_chain.get('oems', []))} tiers={tier_counts}")

            yield f"data: {json.dumps({'type': 'file_ready', 'file_id': file_id, 'filename': filename, 'index': i, 'total': total})}\n\n"

        yield f"data: {json.dumps({'type': 'done', 'total': total})}\n\n"

    return Response(generate(), mimetype="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
    })


@app.route('/dev/multi_single_export/download/<file_id>')
def multi_export_download(file_id):
    with _export_files_lock:
        entry = _export_files.pop(file_id, None)

    if not entry:
        return jsonify({"error": "File not found or already downloaded"}), 404

    path = entry["path"]
    filename = entry["filename"]

    response = send_file(
        path,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=filename
    )

    @response.call_on_close
    def cleanup():
        try:
            os.remove(path)
        except OSError:
            pass

    return response
    

@app.route('/api/map_all')
def map_all():
    collected_oems = []
    collected_tiers = {}
    providers = available_providers()
    product = request.args.get("product", "").strip()
    depth    = max(0, min(int(request.args.get("depth", 2)), 3))
    #provider = request.args.get("provider","gemini").strip().lower()

    if not product:
        return jsonify({"error":"product required"}), 400
    queue, done = [], threading.Event()

    def run():
        try:
            queue.append({"type": "status", "message": "⚡Beginning series of pipeline runs."})
            for provider in providers:
                if not provider.get("configured"):
                    queue.append({"type": "status", "message": f"✅ Skipping {provider.get("name")} due to lack of API key"})
                else:
                    provider_id = provider.get("id")
                    queue.append({"type": "provider_start", "provider": provider["id"], "name": provider["name"], "message": f"Starting Pipeline with {provider['name']}" });
                    run_pipeline(product, depth, provider_id, queue, collected_oems, collected_tiers)
                    queue.append({"type": "provider_done", "provider": provider["id"], "name": provider["name"], "message": f"{provider['name']} pipeline complete."})   
            queue.append({"type": "status", "message": "✅ Completed series of pipeline executions."})
        except Exception as e: queue.append({"type":"error","message":str(e)})
        finally: done.set()
            

    threading.Thread(target=run, daemon=True).start()
 
    def generate():
        sent = 0
        # safer logic to guard against done being set while items are still queued
        while not done.is_set() or sent < len(queue):
            while sent < len(queue):
                yield sse(queue[sent]); sent += 1
            time.sleep(0.1)
        yield sse({"type": "stream_end"})

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})
    


@app.route("/api/providers")
def get_providers():
    return jsonify(available_providers())

@app.route("/api/map")
def map_suppliers():
    product  = request.args.get("product","").strip()
    depth    = max(0, min(int(request.args.get("depth", 2)), 3))
    provider = request.args.get("provider","gemini").strip().lower()

    if not product:
        return jsonify({"error":"product required"}), 400
    if provider not in ("anthropic","openai","gemini","deepseek"):
        return jsonify({"error":f"Unknown provider: {provider}"}), 400

    queue, done = [], threading.Event()

    def run():
        try: run_pipeline(product, depth, provider, queue, [], {})
        except Exception as e: queue.append({"type":"error","message":str(e)})
        finally: done.set()

    threading.Thread(target=run, daemon=True).start()

    def generate():
        sent = 0
        # while not done.is_set() or sent < len(queue): changed while loop condition to test async threads
        while not done.is_set() or sent < len(queue): 
            while sent < len(queue):
                yield sse(queue[sent]); sent += 1
            time.sleep(0.1)
        yield sse({"type":"stream_end"})

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})


# ── Export routes ─────────────────────────────────────────────────────────────

import subprocess
import tempfile
import shutil
from pathlib import Path

REPORT_JS = Path(__file__).parent / "generate_report.js"

def find_node():
    """Find the node executable on Mac/Linux, checking common install paths."""
    # Check PATH first
    node = shutil.which('node')
    if node:
        return node
    # Common Mac paths (Homebrew, nvm, official installer)
    candidates = [
        '/usr/local/bin/node',
        '/opt/homebrew/bin/node',
        '/usr/bin/node',
        str(Path.home() / '.nvm/versions/node'),  # nvm (finds latest)
    ]
    for path in candidates:
        p = Path(path)
        if p.exists() and p.is_file():
            return str(p)
        # nvm directory — find latest version
        if p.is_dir():
            versions = sorted(p.iterdir(), reverse=True)
            for v in versions:
                node_bin = v / 'bin' / 'node'
                if node_bin.exists():
                    return str(node_bin)
    return None


@app.route("/api/export/docx", methods=["POST"])
def export_docx():
    """Generate Word document using pure Python (python-docx). No Node.js needed."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io, datetime

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        product = data.get('product', {})
        tiers   = data.get('tiers', {})
        summary = data.get('summary', '')

        TIER_COLORS = [
            RGBColor(0x2B, 0x7F, 0xCC),
            RGBColor(0xD4, 0xAC, 0x0D),
            RGBColor(0x7D, 0x3C, 0x98),
            RGBColor(0xCA, 0x6F, 0x1E),
        ]
        TIER_HEX = ['2B7FCC','D4AC0D','7D3C98','CA6F1E']

        def set_cell_bg(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  hex_color)
            tcPr.append(shd)

        def add_cell_text(cell, text, bold=False, size=10, color=None, italic=False):
            para = cell.paragraphs[0]
            run  = para.add_run(str(text or '—'))
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # ── Title ──
        title = doc.add_heading('Supply Chain Intelligence Report', 0)
        title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        doc.add_paragraph(
            f"Generated: {datetime.date.today().strftime('%B %d, %Y')}"
        ).runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()

        # ── Product Overview ──
        doc.add_heading('1. Product Overview', level=1)
        prod_rows = [
            ('Product Name',     product.get('product_name','')),
            ('Category',         product.get('product_category','')),
            ('Industry',         product.get('industry','')),
            ('OEM Manufacturer', product.get('oem_manufacturer','')),
            ('OEM Country',      product.get('oem_country','')),
            ('Key Components',   ', '.join(product.get('key_components',[]) or [])),
            ('Description',      product.get('description','')),
        ]
        tbl = doc.add_table(rows=len(prod_rows), cols=2)
        tbl.style = 'Table Grid'
        col_w = [Inches(2), Inches(4.5)]
        for i, (label, value) in enumerate(prod_rows):
            row = tbl.rows[i]
            row.cells[0].width = col_w[0]
            row.cells[1].width = col_w[1]
            add_cell_text(row.cells[0], label, bold=True, size=10)
            add_cell_text(row.cells[1], value, size=10)
            bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
            set_cell_bg(row.cells[0], bg)
            set_cell_bg(row.cells[1], bg)
        doc.add_paragraph()

        # ── OEM Manufacturers ──
        oems = data.get('oems', [])
        if oems:
            doc.add_heading('2. OEM Manufacturers', level=1)
            oem_rows = [('Company','Country','Role','Market Share','Notes')]
            for oem in oems:
                oem_rows.append((
                    oem.get('company_name',''),
                    oem.get('country',''),
                    oem.get('role',''),
                    oem.get('market_share',''),
                    oem.get('notes',''),
                ))
            otbl = doc.add_table(rows=len(oem_rows), cols=5)
            otbl.style = 'Table Grid'
            owidths = [Inches(1.5), Inches(1.0), Inches(1.4), Inches(0.9), Inches(1.7)]
            for ri, row_data in enumerate(oem_rows):
                row = otbl.rows[ri]
                for ci, (val, w) in enumerate(zip(row_data, owidths)):
                    row.cells[ci].width = w
                    is_header = ri == 0
                    add_cell_text(row.cells[ci], val, bold=is_header, size=9,
                                  color=RGBColor(0xFF,0xFF,0xFF) if is_header else None)
                    if is_header:
                        set_cell_bg(row.cells[ci], '1A6B3C')
                    else:
                        set_cell_bg(row.cells[ci], 'F0F8F0' if ri%2==0 else 'FFFFFF')
            doc.add_paragraph()

        # ── Executive Summary ──
        if summary:
            doc.add_heading('2. Executive Summary', level=1)
            p = doc.add_paragraph(summary)
            p.runs[0].font.size = Pt(11)
            doc.add_paragraph()

        # ── Supplier Tiers ──
        doc.add_heading('3. Supplier Tiers', level=1)
        tier_names = ['Tier 1','Tier 2','Tier 3','Tier 4']

        for ti, (tier_key, suppliers) in enumerate(tiers.items()):
            color     = TIER_COLORS[ti] if ti < len(TIER_COLORS) else RGBColor(0x80,0x80,0x80)
            hex_color = TIER_HEX[ti]    if ti < len(TIER_HEX)    else '808080'
            tname     = tier_names[ti]  if ti < len(tier_names)  else f'Tier {ti+1}'

            h = doc.add_heading(f'{tname} Suppliers ({len(suppliers)})', level=2)
            h.runs[0].font.color.rgb = color

            if not suppliers:
                doc.add_paragraph('No suppliers identified.')
                continue

            headers = ['Company','Country','Supplies To','Components Supplied','Confidence']
            tbl2 = doc.add_table(rows=len(suppliers)+1, cols=5)
            tbl2.style = 'Table Grid'
            widths = [Inches(1.6), Inches(1.0), Inches(1.4), Inches(1.9), Inches(0.7)]

            # Header row
            hrow = tbl2.rows[0]
            for ci, (hdr, w) in enumerate(zip(headers, widths)):
                hrow.cells[ci].width = w
                add_cell_text(hrow.cells[ci], hdr, bold=True, size=9,
                              color=RGBColor(0xFF,0xFF,0xFF))
                set_cell_bg(hrow.cells[ci], hex_color)

            # Data rows
            for ri, s in enumerate(suppliers):
                drow = tbl2.rows[ri+1]
                vals = [
                    s.get('company_name',''),
                    s.get('country',''),
                    s.get('supplies_to',''),
                    ', '.join((s.get('components_supplied') or [])[:3]),
                    s.get('confidence',''),
                ]
                bg = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
                for ci, (val, w) in enumerate(zip(vals, widths)):
                    drow.cells[ci].width = w
                    add_cell_text(drow.cells[ci], val, bold=(ci==0), size=9)
                    set_cell_bg(drow.cells[ci], bg)
            doc.add_paragraph()

        # ── Statistics ──
        doc.add_heading('4. Statistics', level=1)
        total     = sum(len(v) for v in tiers.values())
        countries = len(set(s.get('country','') for arr in tiers.values()
                            for s in arr if s.get('country')))
        stats = [
            ('Total Suppliers Identified', str(total)),
            ('Countries Represented',      str(countries)),
            ('Tier Depth Mapped',          str(len(tiers))),
            ('Industry',                   product.get('industry','—')),
        ]
        stbl = doc.add_table(rows=len(stats), cols=2)
        stbl.style = 'Table Grid'
        for i, (label, value) in enumerate(stats):
            stbl.rows[i].cells[0].width = Inches(2.5)
            stbl.rows[i].cells[1].width = Inches(4)
            add_cell_text(stbl.rows[i].cells[0], label, bold=True, size=10)
            add_cell_text(stbl.rows[i].cells[1], value, size=10)
            bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
            set_cell_bg(stbl.rows[i].cells[0], bg)
            set_cell_bg(stbl.rows[i].cells[1], bg)

        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        product_name = product.get('product_name', 'supply_chain')
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', product_name)[:40]

        return Response(
            buf.read(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{safe_name}_supply_chain.docx"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

        return jsonify({"error": str(e)}), 500


@app.route("/api/export/pdf", methods=["POST"])
def export_pdf():
    """Generate and return a PDF from supply chain data using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        import io

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        product = data.get('product', {})
        tiers = data.get('tiers', {})
        summary = data.get('summary', '')

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=inch, leftMargin=inch,
                                topMargin=inch, bottomMargin=inch)

        styles = getSampleStyleSheet()
        DARK = colors.HexColor('#1A1A2E')
        TIER_COLORS_PDF = [
            colors.HexColor('#2B7FCC'),
            colors.HexColor('#D4AC0D'),
            colors.HexColor('#7D3C98'),
            colors.HexColor('#CA6F1E'),
        ]

        title_style  = ParagraphStyle('Title2',  parent=styles['Title'],
                                       fontSize=22, textColor=DARK, spaceAfter=6, fontName='Helvetica-Bold')
        h2_style     = ParagraphStyle('H2',       parent=styles['Heading2'],
                                       fontSize=14, textColor=DARK, spaceBefore=18, spaceAfter=6, fontName='Helvetica-Bold')
        h3_style     = ParagraphStyle('H3',       parent=styles['Heading3'],
                                       fontSize=12, spaceBefore=12, spaceAfter=4, fontName='Helvetica-Bold')
        body_style   = ParagraphStyle('Body2',    parent=styles['Normal'],
                                       fontSize=10, leading=16, spaceAfter=6)
        small_style  = ParagraphStyle('Small',    parent=styles['Normal'],
                                       fontSize=8,  textColor=colors.HexColor('#666666'))

        story = []

        # Title
        story.append(Paragraph('Supply Chain Intelligence Report', title_style))
        story.append(Paragraph(f"Generated: {__import__('datetime').date.today().strftime('%B %d, %Y')}", small_style))
        story.append(Spacer(1, 16))

        # Product overview
        story.append(Paragraph('1. Product Overview', h2_style))
        prod_data = [
            ['Product Name',     product.get('product_name','—')],
            ['Category',         product.get('product_category','—')],
            ['Industry',         product.get('industry','—')],
            ['OEM Manufacturer', product.get('oem_manufacturer','—')],
            ['OEM Country',      product.get('oem_country','—')],
            ['Key Components',   ', '.join(product.get('key_components',[]) or []) or '—'],
        ]
        prod_table = Table(prod_data, colWidths=[2*inch, 4.5*inch])
        prod_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#F0F0F0')),
            ('FONTNAME',   (0,0), (0,-1), 'Helvetica-Bold'),
            ('FONTSIZE',   (0,0), (-1,-1), 9),
            ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
            ('ROWBACKGROUNDS', (0,0), (-1,-1), [colors.HexColor('#F8F8F8'), colors.white]),
            ('VALIGN',     (0,0), (-1,-1), 'TOP'),
            ('PADDING',    (0,0), (-1,-1), 6),
        ]))
        story.append(prod_table)
        story.append(Spacer(1, 12))

        # OEM Manufacturers
        oems = data.get('oems', [])
        if oems:
            story.append(Paragraph('2. OEM Manufacturers', h2_style))
            oem_header = ['Company','Country','Role','Market Share','Notes']
            oem_rows_data = [oem_header] + [
                [o.get('company_name',''), o.get('country',''), o.get('role',''),
                 o.get('market_share',''), o.get('notes','')]
                for o in oems
            ]
            oem_col_w = [1.4*inch, 0.9*inch, 1.2*inch, 0.8*inch, 2.2*inch]
            ot = Table(oem_rows_data, colWidths=oem_col_w, repeatRows=1)
            ot.setStyle(TableStyle([
                ('BACKGROUND',  (0,0), (-1,0), colors.HexColor('#1A6B3C')),
                ('TEXTCOLOR',   (0,0), (-1,0), colors.white),
                ('FONTNAME',    (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE',    (0,0), (-1,-1), 8),
                ('GRID',        (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F0F8F0'), colors.white]),
                ('VALIGN',      (0,0), (-1,-1), 'TOP'),
                ('PADDING',     (0,0), (-1,-1), 5),
            ]))
            story.append(ot)
            story.append(Spacer(1, 12))

        # Summary
        if summary:
            story.append(Paragraph('2. Executive Summary', h2_style))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 12))

        # Tiers
        story.append(Paragraph('3. Supplier Tiers', h2_style))
        tier_names = ['Tier 1','Tier 2','Tier 3','Tier 4']
        for ti, (tier_key, suppliers) in enumerate(tiers.items()):
            color = TIER_COLORS_PDF[ti] if ti < len(TIER_COLORS_PDF) else colors.grey
            tname = tier_names[ti] if ti < len(tier_names) else f'Tier {ti+1}'
            story.append(Paragraph(f'{tname} Suppliers ({len(suppliers)})', h3_style))
            if not suppliers:
                story.append(Paragraph('No suppliers identified.', small_style))
                continue
            header = ['Company', 'Country', 'Supplies To', 'Components', 'Confidence']
            rows = [header] + [
                [s.get('company_name','—'), s.get('country','—'), s.get('supplies_to','—'),
                 ', '.join((s.get('components_supplied') or [])[:2]),
                 s.get('confidence','—')]
                for s in suppliers
            ]
            col_w = [1.5*inch, 1*inch, 1.5*inch, 2*inch, 0.8*inch]
            t = Table(rows, colWidths=col_w, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND',  (0,0), (-1,0), color),
                ('TEXTCOLOR',   (0,0), (-1,0), colors.white),
                ('FONTNAME',    (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE',    (0,0), (-1,-1), 8),
                ('GRID',        (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#F8F8F8'), colors.white]),
                ('VALIGN',      (0,0), (-1,-1), 'TOP'),
                ('PADDING',     (0,0), (-1,-1), 5),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))

        # Stats
        story.append(Paragraph('4. Statistics', h2_style))
        total = sum(len(v) for v in tiers.values())
        countries = len(set(s.get('country','') for arr in tiers.values() for s in arr if s.get('country')))
        stats_data = [
            ['Total Suppliers',    str(total)],
            ['Countries',          str(countries)],
            ['Tier Depth',         str(len(tiers))],
            ['Industry',           product.get('industry','—')],
        ]
        st = Table(stats_data, colWidths=[3*inch, 3.5*inch])
        st.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#F0F0F0')),
            ('FONTNAME',   (0,0), (0,-1), 'Helvetica-Bold'),
            ('FONTSIZE',   (0,0), (-1,-1), 9),
            ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
            ('ROWBACKGROUNDS', (0,0), (-1,-1), [colors.HexColor('#F8F8F8'), colors.white]),
            ('PADDING',    (0,0), (-1,-1), 6),
        ]))
        story.append(st)

        doc.build(story)
        pdf_bytes = buffer.getvalue()

        product_name = product.get('product_name', 'supply_chain')
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', product_name)[:40]

        return Response(
            pdf_bytes,
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{safe_name}_supply_chain.pdf"'}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/health")
def health():
    return jsonify({"status":"ok","providers":available_providers()})

if __name__ == "__main__":
    print("🚀 Supplier Chain Mapper — multi-provider")
    for p in available_providers():
        status = "✅ ready" if p["configured"] else "❌ no key"
        print(f"   {p['name']:25s} {status:12s} search: {p['search']}")
    print("   Running on http://localhost:5000")
    app.run(debug=False, port=5000, threaded=True)
